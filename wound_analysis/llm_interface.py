from typing import Dict, List, Optional
from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline, AutoModelForMaskedLM
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from datetime import datetime
import logging
import torch
import os
import pathlib
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

class WoundAnalysisLLM:
    MODEL_CATEGORIES = {
        "huggingface": {
            "medalpaca-7b": "medalpaca/medalpaca-7b",
            "BioMistral-7B": "BioMistral/BioMistral-7B",
            "ClinicalCamel-70B": "wanglab/ClinicalCamel-70B",
        },
        "ai-verde": {
            # "llama-3-90b-vision": "llama-3-2-90b-vision-instruct-quantized",
            "llama-3.3-70b-fp8": "js2/Llama-3.3-70B-Instruct-FP8-Dynamic",
            "meta-llama-3.1-70b": "Meta-Llama-3.1-70B-Instruct-quantized",  # default
            "deepseek-r1": "js2/DeepSeek-R1",
            "llama-3.2-11b-vision": "Llama-3.2-11B-Vision-Instruct"
        }
    }

    def __init__(self, platform: str = "ai-verde", model_name: str = "llama-3.3-70b-fp8"):
        """
        Initialize the LLM interface with HuggingFace models or AI Verde.
        Args:
            platform: The platform to use ("huggingface" or "ai-verde")
            model_name: Name of the model to use within the selected platform
        """
        if platform not in self.MODEL_CATEGORIES:
            raise ValueError(f"Platform {platform} not supported. Choose from: {list(self.MODEL_CATEGORIES.keys())}")

        if model_name not in self.MODEL_CATEGORIES[platform]:
            raise ValueError(f"Model {model_name} not supported for platform {platform}. Choose from: {list(self.MODEL_CATEGORIES[platform].keys())}")

        self.platform = platform
        self.model_name = model_name
        self.model_path = self.MODEL_CATEGORIES[platform][model_name]
        self.model = None  # Initialize as None, will be loaded on first use

    def _load_model(self):
        """Lazy loading of the model to avoid Streamlit file watcher issues"""
        if self.model is not None:
            return

        try:
            if self.platform == "ai-verde":
                if not os.getenv("OPENAI_API_KEY"):
                    raise ValueError("OPENAI_API_KEY environment variable must be set for AI Verde")

                self.model = ChatOpenAI(
                    model=self.model_path,
                    base_url="https://llm-api.cyverse.ai"
                )
                logger.info(f"Successfully loaded AI Verde model {self.model_name}")
            else:
                device = "cuda" if torch.cuda.is_available() else "cpu"

                if self.model_name == "clinical-bert":
                    self.model = pipeline(
                        "fill-mask",
                        model=self.model_path,
                        device=device
                    )
                else:
                    self.model = pipeline(
                        "text-generation",
                        model=self.model_path,
                        device=device,
                        max_new_tokens=512
                    )
                logger.info(f"Successfully loaded model {self.model_name} on {device}")

        except Exception as e:
            logger.error(f"Error loading model {self.model_name}: {str(e)}")
            raise

    @classmethod
    def get_available_platforms(cls) -> List[str]:
        """Get list of supported platforms."""
        return list(cls.MODEL_CATEGORIES.keys())

    @classmethod
    def get_available_models(cls, platform: str) -> List[str]:
        """Get list of supported models for a specific platform."""
        if platform not in cls.MODEL_CATEGORIES:
            raise ValueError(f"Platform {platform} not supported")
        return list(cls.MODEL_CATEGORIES[platform].keys())

    def _format_prompt(self, patient_data: Dict) -> str:
        """Format patient data into a prompt for the LLM."""
        metadata = patient_data['patient_metadata']
        visits = patient_data['visits']

        # Create a more detailed patient profile
        prompt = (
            "Task: Analyze wound healing progression and provide clinical recommendations.\n\n"
            "Patient Profile:\n"
            f"- {metadata.get('age', 'Unknown')}y/o {metadata.get('sex', 'Unknown')}\n"
            f"- Race/Ethnicity: {metadata.get('race', 'Unknown')}, {metadata.get('ethnicity', 'Unknown')}\n"
            f"- BMI: {metadata.get('bmi', 'Unknown')}\n"
            f"- Study Cohort: {metadata.get('study_cohort', 'Unknown')}\n"
            f"- Smoking: {metadata.get('smoking_status', 'None')}"
        )

        # Add diabetes information
        diabetes_info = metadata.get('diabetes', {})
        if diabetes_info:
            prompt += (
                f"\n- Diabetes: {diabetes_info.get('status', 'None')}\n"
                f"- HbA1c: {diabetes_info.get('hemoglobin_a1c', 'Unknown')}%"
            )

        # Add medical history if available
        medical_history = metadata.get('medical_history', {})
        if medical_history:
            conditions = [cond for cond, val in medical_history.items() if val and val != 'None']
            if conditions:
                prompt += f"\n- Medical History: {', '.join(conditions)}"

        prompt += "\n\nVisit History:\n"

        # Format visit information
        for visit in visits:
            visit_date = visit.get('visit_date', 'Unknown')
            measurements = visit.get('wound_measurements', {})
            wound_info = visit.get('wound_info', {})
            sensor = visit.get('sensor_data', {})

            prompt += (
                f"\nDate: {visit_date}\n"
                f"Location: {wound_info.get('location', 'Unknown')}\n"
                f"Type: {wound_info.get('type', 'Unknown')}\n"
                f"Size: {measurements.get('length')}cm x {measurements.get('width')}cm x {measurements.get('depth')}cm\n"
                f"Area: {measurements.get('area')}cm²\n"
            )

            # Add granulation information
            granulation = wound_info.get('granulation', {})
            if granulation:
                prompt += f"Tissue: {granulation.get('quality', 'Unknown')}, coverage {granulation.get('coverage', 'Unknown')}\n"

            # Add exudate information
            exudate = wound_info.get('exudate', {})
            if exudate:
                prompt += (
                    f"Exudate: {exudate.get('volume', 'Unknown')} volume, "
                    f"viscosity {exudate.get('viscosity', 'Unknown')}, "
                    f"type {exudate.get('type', 'Unknown')}\n"
                )

            # Add sensor data
            if sensor:
                temp = sensor.get('temperature', {})
                impedance = sensor.get('impedance', {})
                high_freq_imp = impedance.get('high_frequency', {})
                center_freq_imp = impedance.get('center_frequency', {})
                low_freq_imp = impedance.get('low_frequency', {})

                prompt += (
                    f"Measurements:\n"
                    f"- O₂: {sensor.get('oxygenation')}%\n"
                    f"- Temperature: center {temp.get('center')}°F, edge {temp.get('edge')}°F, peri-wound {temp.get('peri')}°F\n"
                    f"- Hemoglobin: {sensor.get('hemoglobin')}, Oxy: {sensor.get('oxyhemoglobin')}, Deoxy: {sensor.get('deoxyhemoglobin')}\n"
                )

                # Add all three frequency impedance measurements
                if high_freq_imp and any(v is not None for v in high_freq_imp.values()):
                    prompt += f"- High Frequency Impedance ({high_freq_imp.get('frequency', '80kHz')}): |Z|={high_freq_imp.get('Z')}, resistance={high_freq_imp.get('resistance')}, capacitance={high_freq_imp.get('capacitance')}\n"

                if center_freq_imp and any(v is not None for v in center_freq_imp.values()):
                    prompt += f"- Center Frequency Impedance ({center_freq_imp.get('frequency', '40kHz')}): |Z|={center_freq_imp.get('Z')}, resistance={center_freq_imp.get('resistance')}, capacitance={center_freq_imp.get('capacitance')}\n"

                if low_freq_imp and any(v is not None for v in low_freq_imp.values()):
                    prompt += f"- Low Frequency Impedance ({low_freq_imp.get('frequency', '100Hz')}): |Z|={low_freq_imp.get('Z')}, resistance={low_freq_imp.get('resistance')}, capacitance={low_freq_imp.get('capacitance')}\n"

            # Add infection and treatment information
            infection = wound_info.get('infection', {})
            if infection.get('status') or infection.get('classification'):
                prompt += f"Infection Status: {infection.get('status', 'None')}, Classification: {infection.get('classification', 'None')}\n"

            if wound_info.get('current_care'):
                prompt += f"Current Care: {wound_info.get('current_care')}\n"

        prompt += (
            "\nPlease provide a comprehensive analysis including:\n"
            "1. Wound healing trajectory (analyze changes in size, exudate, and tissue characteristics)\n"
            "2. Concerning patterns (identify any worrying trends in measurements or characteristics)\n"
            "3. Care recommendations (based on wound type, characteristics, and healing progress)\n"
            "4. Complication risks (assess based on patient profile and wound characteristics)\n"
            "5. Significance of sensor measurements (interpret oxygenation, temperature, and impedance trends)\n"
        )

        return prompt

    def _format_population_prompt(self, population_data: Dict) -> str:
        """Format population-level data into a prompt for the LLM."""

        prompt = (
            "Task: Analyze wound healing patterns and provide comprehensive clinical insights for the entire patient population.\n\n"
            "Population Overview:\n"
        )

        # Add demographic statistics
        demographics = population_data.get('demographics', {})
        prompt += (
            f"Total Patients: {demographics.get('total_patients', 'Unknown')}\n"
            f"Age Distribution: {demographics.get('age_stats', 'Unknown')}\n"
            f"Gender Distribution: {demographics.get('gender_distribution', 'Unknown')}\n"
            f"Race/Ethnicity Distribution: {demographics.get('race_distribution', 'Unknown')}\n"
            f"BMI Distribution: {demographics.get('bmi_stats', 'Unknown')}\n"
            f"Smoking Status Distribution: {demographics.get('smoking_distribution', 'Unknown')}\n"
            f"Diabetes Status Distribution: {demographics.get('diabetes_distribution', 'Unknown')}\n"
            f"Comorbidity Distribution: {demographics.get('comorbidity_distribution', 'Unknown')}\n"
        )

        # Add wound characteristics with more detail
        wound_stats = population_data.get('wound_stats', {})
        prompt += (
            "\nWound Characteristics:\n"
            f"Wound Types Distribution: {wound_stats.get('type_distribution', 'Unknown')}\n"
            f"Location Distribution: {wound_stats.get('location_distribution', 'Unknown')}\n"
            f"Average Initial Size: {wound_stats.get('avg_initial_size', 'Unknown')}\n"
            f"Size Distribution: {wound_stats.get('size_distribution', 'Unknown')}\n"
            f"Average Treatment Duration: {wound_stats.get('avg_treatment_duration', 'Unknown')}\n"
            f"Duration Distribution: {wound_stats.get('duration_distribution', 'Unknown')}\n"
            f"Initial Wound Grade Distribution: {wound_stats.get('initial_grade_distribution', 'Unknown')}\n"
        )

        # Add healing rate analysis with detailed statistics
        healing_stats = population_data.get('healing_stats', {})
        prompt += (
            "\nHealing Rate Analysis:\n"
            f"Overall Healing Rate Distribution: {healing_stats.get('overall_distribution', 'Unknown')}\n"
            f"Average Weekly Healing Rate: {healing_stats.get('avg_weekly_rate', 'Unknown')}\n"
            f"Complete Healing Rate: {healing_stats.get('complete_healing_rate', 'Unknown')}%\n"
            f"Time to 50% Healing: {healing_stats.get('time_to_50_percent', 'Unknown')}\n"
            "\nHealing Rates by Risk Factors:\n"
            f"- Diabetes Impact: {healing_stats.get('diabetes_impact', 'Unknown')}\n"
            f"- Smoking Impact: {healing_stats.get('smoking_impact', 'Unknown')}\n"
            f"- BMI Impact: {healing_stats.get('bmi_impact', 'Unknown')}\n"
            f"- Age Impact: {healing_stats.get('age_impact', 'Unknown')}\n"
            f"- Comorbidity Impact: {healing_stats.get('comorbidity_impact', 'Unknown')}\n"
        )

        # Add treatment outcome analysis
        treatment_stats = population_data.get('treatment_stats', {})
        prompt += (
            "\nTreatment Outcomes:\n"
            f"Treatment Modalities Used: {treatment_stats.get('modalities_used', 'Unknown')}\n"
            f"Success Rates by Treatment: {treatment_stats.get('success_rates', 'Unknown')}\n"
            f"Average Time to Response: {treatment_stats.get('time_to_response', 'Unknown')}\n"
            f"Treatment Adherence Rates: {treatment_stats.get('adherence_rates', 'Unknown')}\n"
            f"Complication Rates: {treatment_stats.get('complication_rates', 'Unknown')}\n"
        )

        # Add exudate analysis with trends
        exudate_stats = population_data.get('exudate_stats', {})
        prompt += (
            "\nExudate Analysis:\n"
            f"Volume Distribution: {exudate_stats.get('volume_distribution', 'Unknown')}\n"
            f"Type Distribution: {exudate_stats.get('type_distribution', 'Unknown')}\n"
            f"Viscosity Distribution: {exudate_stats.get('viscosity_distribution', 'Unknown')}\n"
            f"Volume Trends Over Time: {exudate_stats.get('volume_trends', 'Unknown')}\n"
            f"Correlation with Healing: {exudate_stats.get('healing_correlation', 'Unknown')}\n"
        )

        # Add sensor data statistics with trends and correlations
        sensor_stats = population_data.get('sensor_stats', {})
        if sensor_stats:
            prompt += (
                "\nSensor Measurements Analysis:\n"
                f"Oxygenation Levels: {sensor_stats.get('oxygenation_stats', 'Unknown')}\n"
                f"Oxygenation Trends: {sensor_stats.get('oxygenation_trends', 'Unknown')}\n"
                f"Temperature Patterns: {sensor_stats.get('temperature_stats', 'Unknown')}\n"
                f"Temperature Trends: {sensor_stats.get('temperature_trends', 'Unknown')}\n"
                f"Impedance Measurements: {sensor_stats.get('impedance_stats', 'Unknown')}\n"
                f"Impedance Trends: {sensor_stats.get('impedance_trends', 'Unknown')}\n"
                "\nSensor-Healing Correlations:\n"
                f"- Oxygenation Impact: {sensor_stats.get('oxygenation_impact', 'Unknown')}\n"
                f"- Temperature Impact: {sensor_stats.get('temperature_impact', 'Unknown')}\n"
                f"- Impedance Impact: {sensor_stats.get('impedance_impact', 'Unknown')}\n"
            )

        # Add correlation analysis with detailed insights
        correlations = population_data.get('correlations', {})
        if correlations:
            prompt += (
                "\nMultivariate Analysis:\n"
                "Healing Rate Correlations:\n"
                f"- Patient Factors: {correlations.get('patient_factors', 'Unknown')}\n"
                f"- Wound Characteristics: {correlations.get('wound_characteristics', 'Unknown')}\n"
                f"- Treatment Approaches: {correlations.get('treatment_approaches', 'Unknown')}\n"
                f"- Sensor Measurements: {correlations.get('sensor_measurements', 'Unknown')}\n"
                f"\nPredictive Factors:\n{correlations.get('predictive_factors', 'Unknown')}\n"
                f"Risk Factor Interactions: {correlations.get('risk_interactions', 'Unknown')}\n"
            )

        prompt += (
            "\nPlease provide a comprehensive analysis of the population-level data, including:\n"
            "1. Key Patterns and Trends:\n"
            "   - Identify significant patterns in wound healing across different patient groups\n"
            "   - Analyze temporal trends in healing rates and wound characteristics\n"
            "   - Highlight any unexpected or notable findings\n\n"
            "2. Risk Factor Analysis:\n"
            "   - Evaluate the impact of each risk factor on healing outcomes\n"
            "   - Identify high-risk patient profiles\n"
            "   - Analyze interactions between multiple risk factors\n\n"
            "3. Treatment Effectiveness:\n"
            "   - Compare outcomes across different treatment modalities\n"
            "   - Identify factors associated with treatment success\n"
            "   - Analyze time-to-response patterns\n\n"
            "4. Sensor Data Insights:\n"
            "   - Interpret patterns in oxygenation, temperature, and impedance measurements\n"
            "   - Correlate sensor data with healing outcomes\n"
            "   - Identify predictive indicators\n\n"
            "5. Clinical Recommendations:\n"
            "   - Suggest evidence-based treatment strategies\n"
            "   - Recommend risk mitigation approaches\n"
            "   - Propose monitoring protocols based on risk profiles\n\n"
            "6. Future Directions:\n"
            "   - Identify areas needing additional investigation\n"
            "   - Suggest potential protocol improvements\n"
            "   - Recommend data collection enhancements\n"
        )

        return prompt

    def analyze_patient_data(self, patient_data: Dict) -> str:
        """
        Analyze patient data using the configured model.
        Args:
            patient_data: Processed patient data dictionary
        Returns:
            Analysis results as string
        """
        # Load model if not already loaded
        self._load_model()
        prompt = self._format_prompt(patient_data)

        try:
            if self.platform == "ai-verde":
                # Use AI Verde with system and human messages
                messages = [
                    SystemMessage(content="You are a medical expert specializing in wound care analysis. Analyze the provided wound data and provide clinical recommendations."),
                    HumanMessage(content=prompt)
                ]
                response = self.model.invoke(messages)
                return response.content

            elif self.model_name == "clinical-bert":
                # For BERT, we'll use a simpler approach focused on key aspects
                latest_visit = patient_data['visits'][-1]
                measurements = latest_visit.get('wound_measurements', {})
                wound_info = latest_visit.get('wound_info', {})

                # Create shorter, focused prompts
                healing_prompt = f"The wound healing progress is [MASK]. Size is {measurements.get('length')}x{measurements.get('width')}cm."
                risk_prompt = f"The wound infection risk is [MASK]. Exudate is {wound_info.get('exudate', {}).get('volume', 'Unknown')}."

                # Get predictions for each aspect
                healing_result = self.model(healing_prompt)[0]['token_str']
                risk_result = self.model(risk_prompt)[0]['token_str']

                # Combine results
                analysis = f"""
                Wound Analysis Summary:
                - Healing Progress: {healing_result}
                - Risk Assessment: {risk_result}
                - Latest Measurements: {measurements.get('length')}cm x {measurements.get('width')}cm
                - Area: {measurements.get('area')}cm²
                """
                return analysis
            else:
                # Standard text generation for other models
                response = self.model(
                    prompt,
                    do_sample=True,
                    temperature=0.3,
                    max_new_tokens=512,
                    num_return_sequences=1
                )

                # Extract and clean the generated text
                generated_text = response[0]['generated_text']
                # Remove the prompt from the response if it's included
                if generated_text.startswith(prompt):
                    generated_text = generated_text[len(prompt):].strip()

                return generated_text

        except Exception as e:
            logger.error(f"Error during analysis: {str(e)}")
            raise

    def analyze_population_data(self, population_data: Dict) -> str:
        """
        Analyze population-level data using the configured model.
        Args:
            population_data: Dictionary containing aggregated population statistics and analysis
        Returns:
            Analysis results as string
        """
        self._load_model()

        try:
            prompt = self._format_population_prompt(population_data)

            if self.platform == "ai-verde":
                system_prompt = (
                    "You are a senior wound care specialist and data scientist analyzing population-level wound healing data. "
                    "Your expertise spans clinical wound care, biostatistics, and medical data analysis. Your task is to:\n\n"
                    "1. Analyze complex wound healing patterns across diverse patient populations\n"
                    "2. Identify clinically significant trends and correlations\n"
                    "3. Evaluate the effectiveness of different treatment approaches\n"
                    "4. Provide evidence-based recommendations for improving wound care outcomes\n\n"
                    "Focus on actionable insights that have practical clinical applications. Support your analysis with "
                    "specific data points and statistics when available. Consider both statistical significance and "
                    "clinical relevance in your interpretations. When making recommendations, consider implementation "
                    "feasibility and potential resource constraints."
                )

                messages = [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=prompt)
                ]
                response = self.model.invoke(messages)
                return response.content
            else:
                # For HuggingFace models, we'll use a simplified version of the prompt
                combined_prompt = (
                    "You are a wound care expert analyzing population-level data. "
                    "Provide a detailed analysis focusing on patterns, risk factors, "
                    "and evidence-based recommendations.\n\n" + prompt
                )

                response = self.model(
                    combined_prompt,
                    do_sample=True,
                    temperature=0.3,
                    max_new_tokens=1024,  # Increased for more detailed analysis
                    num_return_sequences=1
                )

                if isinstance(response, list):
                    generated_text = response[0]['generated_text']
                else:
                    generated_text = response['generated_text']

                # Clean up the response
                analysis = generated_text.replace(combined_prompt, '').strip()

                # Format the analysis with clear sections
                sections = [
                    "Key Patterns and Trends",
                    "Risk Factor Analysis",
                    "Treatment Effectiveness",
                    "Sensor Data Insights",
                    "Clinical Recommendations",
                    "Future Directions"
                ]

                formatted_analysis = ""
                current_section = ""
                for line in analysis.split('\n'):
                    # Check if line starts a new section
                    is_section = False
                    for section in sections:
                        if section.lower() in line.lower():
                            current_section = section
                            formatted_analysis += f"\n## {section}\n"
                            is_section = True
                            break

                    if not is_section and line.strip():
                        if not current_section:
                            current_section = "Analysis"
                            formatted_analysis += "\n## General Analysis\n"
                        formatted_analysis += line + "\n"

                return formatted_analysis.strip()

        except Exception as e:
            logger.error(f"Error analyzing population data: {str(e)}")
            raise


def create_and_save_report(patient_metadata: dict, analysis_results: str) -> str:
    """Create and save the analysis report as a Word document."""
    doc = Document()
    report_path = format_word_document(doc=doc, analysis_results=analysis_results, patient_metadata=patient_metadata)
    return report_path

def download_word_report(st, report_path: str):
    """Create a download link for the Word report."""
    try:
        with open(report_path, 'rb') as f:
            bytes_data = f.read()
            st.download_button(
                label="Download Full Report (DOCX)",
                data=bytes_data,
                file_name=os.path.basename(report_path),
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    except Exception as e:
        st.error(f"Error preparing report download: {str(e)}")

def format_word_document(doc: Document, analysis_results: str, patient_metadata: dict=None, report_path: str = None) -> str:
    """
    Format the analysis results in a professional Word document layout.
    Returns the path to the saved document.
    """
    # Add title
    title = doc.add_heading('Wound Care Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if patient_metadata is not None:
        # Add patient information section
        doc.add_heading('Patient Information', level=1)
        patient_info = doc.add_paragraph()
        patient_info.add_run('Patient Demographics:\n').bold = True
        patient_info.add_run(f"Age: {patient_metadata.get('age', 'Unknown')} years\n")
        patient_info.add_run(f"Sex: {patient_metadata.get('sex', 'Unknown')}\n")
        patient_info.add_run(f"BMI: {patient_metadata.get('bmi', 'Unknown')}\n")

        # Add diabetes information
        diabetes_info = doc.add_paragraph()
        diabetes_info.add_run('Diabetes Status:\n').bold = True
        if 'diabetes' in patient_metadata:
            diabetes_info.add_run(f"Type: {patient_metadata['diabetes'].get('status', 'Unknown')}\n")
            diabetes_info.add_run(f"HbA1c: {patient_metadata['diabetes'].get('hemoglobin_a1c', 'Unknown')}%\n")

    # Add analysis section
    doc.add_heading('Analysis Results', level=1)

    # Split analysis into sections and format them
    sections = analysis_results.split('\n\n')
    for section in sections:
        if section.strip():
            if '**' in section:  # Handle markdown-style headers
                # Convert markdown headers to proper formatting
                section = section.replace('**', '')
                p = doc.add_paragraph()
                p.add_run(section.strip()).bold = True
            else:
                # Handle bullet points
                if section.strip().startswith('- ') or section.strip().startswith('* '):
                    p = doc.add_paragraph(section.strip()[2:], style='List Bullet')
                else:
                    p = doc.add_paragraph(section.strip())

    # Add footer with timestamp
    doc.add_paragraph(f"\nReport generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Save the document
    if report_path is None:
        # Create logs directory if it doesn't exist
        log_dir = pathlib.Path(__file__).parent / 'logs'
        log_dir.mkdir(exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_path = log_dir / f'wound_analysis_{timestamp}.docx'

    doc.save(report_path)
    return str(report_path)

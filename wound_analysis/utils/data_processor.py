import logging
import os
import pathlib
import re
import traceback
from dataclasses import dataclass, field
from datetime import datetime
from typing import Dict, Optional

import numpy as np
import math
from scipy import stats
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .column_schema import DataColumns

try:
    from ..dashboard import debug_log
except ImportError:
    # Fallback if we can't import from dashboard
    def debug_log(message):
        print(message)
        try:
            with open('/app/debug.log', 'a') as f:
                f.write(f"{message}\n")
        except:
            pass

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class WoundDataProcessor:
	def __init__(self, csv_dataset_path: pathlib.Path=None, impedance_freq_sweep_path: pathlib.Path=None, df: pd.DataFrame= None):

		self.columns = DataColumns()
		self.impedance_freq_sweep_path = impedance_freq_sweep_path

		self.df = pd.read_csv( csv_dataset_path ) if df is None else df

		# Clean column names
		self.df.columns = self.df.columns.str.strip()

	def get_patient_visits(self, record_id: int) -> Dict:
		"""
			Retrieves all visit data for a specific patient.

			This method fetches all visits for a patient identified by their record ID, processes
			the data for each non-skipped visit, and includes wound information.

			Args:
				record_id (int): The unique identifier for the patient.

			Returns:
				Dict: A dictionary containing:
					- 'patient_metadata': Basic patient information extracted from their first visit
					- 'visits': List of dictionaries containing data for each valid visit, including wound information

			Raises:
				ValueError: If no records are found for the specified patient
				Exception: If any error occurs during data processing
		"""

		try:
			patient_id_col    = self.columns.patient_identifiers.record_id
			skipped_visit_col = self.columns.visit_info.skipped_visit

			patient_data = self.df[self.df[patient_id_col] == record_id]
			if patient_data.empty:
				raise ValueError(f"No measurements found for patient {record_id}")

			# Get metaa from first visit
			first_visit = patient_data.iloc[0]
			metadata = self._extract_patient_metadata(first_visit)

			visits_data = []
			for _, visit in patient_data.iterrows():
				if pd.isna(visit.get(skipped_visit_col)) or visit[skipped_visit_col] != 'Yes':
					visit_data = self._process_visit_data(visit=visit, record_id=record_id)
					if visit_data:
						wound_info = self._get_wound_info(visit)
						visit_data['wound_info'] = wound_info
						visits_data.append(visit_data)

			return {
				'patient_metadata': metadata,
				'visits': visits_data
			}
		except Exception as e:
			logger.error(f"Error processing patient {record_id}: {str(e)}")
			raise

	def get_population_statistics(self) -> Dict:
		"""
			Gather comprehensive population-level statistics for LLM analysis.

			This method aggregates data from all patients and visits to provide a holistic view
			of the wound care dataset. It includes information from various aspects of wound care,
			patient demographics, and treatment outcomes.

			Returns:
				Dict: A comprehensive dictionary containing population statistics including:
					- Summary (total patients, visits, overall improvement rate, etc.)
					- Demographics (ag gender, race, BMI distribution)
					- Wound characteristics (types, locations, initial sizes)
					- Healing progression metrics (healing rates, time to closure)
					- Sensor data analysis (impedance trends, temperature patterns, oxygenation levels)
					- Risk factor analysis (comorbidities, listyle factors)
					- Treatment effectiveness (comparison of different approaches)
					- Temporal trends (seasonal variations, long-term outcome improvements)

			Note:
				This method processes the entire dataset and may be computationally intensive
				for large datasets.
		"""
		# Get schema columns
		pi   = self.columns.patient_identifiers
		vis  = self.columns.visit_info
		wc   = self.columns.wound_characteristics
		dem  = self.columns.demographics
		ls   = self.columns.lifestyle
		mh   = self.columns.medical_history
		temp = self.columns.temperature
		oxy  = self.columns.oxygenation
		imp  = self.columns.impedance
		ca   = self.columns.clinical_assessment
		hm   = self.columns.healing_metrics

		# Get processed data
		df = self.get_processed_data()
		if df.empty:
			raise ValueError("No data available for population statistics")

		# Calculate derived metrics
		df['BMI Category'] = pd.cut(
			df[dem.bmi],
			bins=[0, 18.5, 24.9, 29.9, float('inf')],
			labels=['Underweight', 'Normal', 'Overweight', 'Obese']
		)

		# Calculate healing metrics
		df['Healing_Color']  = df[hm.healing_rate].apply(lambda x: 'green' if x < 0 else 'red')
		df['Healing_Status'] = df[hm.healing_rate].apply(
			lambda x: 'Improving' if x < 0 else ('Stable' if -5 <= x <= 5 else 'Worsening')
		)

		stats_data = {
			'summary': {
				'total_patients'             : len(df[pi.record_id].unique()),
				'total_visits'               : len(df),
				'avg_visits_per_patient'     : len(df) / len(df[pi.record_id].unique()),
				'overall_improvement_rate'   : (df[hm.healing_rate] < 0).mean() * 100,
				'avg_treatment_duration_days': (df.groupby(pi.record_id)[vis.days_since_first_visit].max().mean()),
				'completion_rate'            : (df['Visit Status'] == 'Completed').mean() * 100 if 'Visit Status' in df.columns else None
			},
			'demographics': {
				'age_stats': {
					'summary': f"Mean: {df[dem.age_at_enrollment].mean():.1f}, Median: {df[dem.age_at_enrollment].median():.1f}",
					'distribution': df[dem.age_at_enrollment].value_counts().to_dict(),
					'age_groups': pd.cut(df[dem.age_at_enrollment],
						bins=[0, 30, 50, 70, float('inf')],
						labels=['<30', '30-50', '50-70', '>70']).value_counts().to_dict()
				},
				'gender_distribution'   : df[dem.sex].value_counts().to_dict(),
				'race_distribution'     : df[dem.race].value_counts().to_dict(),
				'ethnicity_distribution': df[dem.ethnicity].value_counts().to_dict(),
				'bmi_stats': {
					'summary': f"Mean: {df[dem.bmi].mean():.1f}, Range: {df[dem.bmi].min():.1f}-{df[dem.bmi].max():.1f}",
					'distribution'     : df['BMI Category'].value_counts().to_dict(),
					'by_healing_status': df.groupby('BMI Category')[hm.healing_rate].agg(['mean', 'count']).to_dict()
				}
			},
			'risk_factors': {
				'primary_conditions': {
					'diabetes': {
						'distribution'  : df[mh.diabetes].value_counts().to_dict(),
						'healing_impact': df.groupby(mh.diabetes)[hm.healing_rate].agg(['mean', 'std', 'count']).to_dict()
					},
					'smoking': {
						'distribution'  : df[ls.smoking_status].value_counts().to_dict(),
						'healing_impact': df.groupby(ls.smoking_status)[hm.healing_rate].agg(['mean', 'std', 'count']).to_dict()
					}
				},
				'comorbidity_analysis': {
					'diabetes_smoking': df.groupby([mh.diabetes, ls.smoking_status])[hm.healing_rate].agg(['mean', 'count']).to_dict(),
					'diabetes_bmi'    : df.groupby([mh.diabetes, 'BMI Category'])[hm.healing_rate].agg(['mean', 'count']).to_dict()
				}
			},
			'wound_characteristics': {
				'type_distribution': {
					'overall': df[wc.wound_type].value_counts().to_dict(),
					'by_healing_status': df.groupby([wc.wound_type, 'Healing_Status']).size().to_dict()
				},
				'location_analysis': {
					'distribution': df[wc.wound_location].value_counts().to_dict(),
					'healing_by_location': df.groupby(wc.wound_location)[hm.healing_rate].mean().to_dict()
				},
				'size_progression': {
					'initial_vs_final': {
						'area': {
							'initial': df.groupby(pi.record_id)[wc.wound_area].first().agg(['mean', 'median', 'std']).to_dict(),
							'final'  : df.groupby(pi.record_id)[wc.wound_area].last().agg(['mean', 'median', 'std']).to_dict(),
							'percent_change': ((df.groupby(pi.record_id)[wc.wound_area].last() -
												df.groupby(pi.record_id)[wc.wound_area].first()) /
												df.groupby(pi.record_id)[wc.wound_area].first() * 100).mean()
						}
					},
					'healing_by_initial_size': {
						'small' : df[df[wc.wound_area] < df[wc.wound_area].quantile(0.33)][hm.healing_rate].mean(),
						'medium': df[(df[wc.wound_area] >= df[wc.wound_area].quantile(0.33)) &
									(df[wc.wound_area] < df[wc.wound_area].quantile(0.67))][hm.healing_rate].mean(),
						'large': df[df[wc.wound_area] >= df[wc.wound_area].quantile(0.67)][hm.healing_rate].mean()
					}
				}
			},
			'healing_progression': {
				'overall_stats': {
					'summary': f"Mean: {df[hm.healing_rate].mean():.1f}%, Median: {df[hm.healing_rate].median():.1f}%",
					'distribution': df['Healing_Status'].value_counts().to_dict(),
					'percentiles' : df[hm.healing_rate].quantile([0.25, 0.5, 0.75]).to_dict()
				},
				'temporal_analysis': {
					'by_visit_number': df.groupby('Visit Number')[hm.healing_rate].agg(['mean', 'std', 'count']).to_dict(),
					'by_treatment_duration': pd.cut(df[vis.days_since_first_visit],
						bins=[0, 30, 90, 180, float('inf')],
						labels=['<30 days', '30-90 days', '90-180 days', '>180 days']
					).value_counts().to_dict()
				}
			},
			'exudate_analysis': {
				'characteristics': {
					'volume': {
						'distribution'       : df[ca.exudate_volume].value_counts().to_dict(),
						'healing_correlation': df.groupby(ca.exudate_volume)[hm.healing_rate].mean().to_dict()
					},
					'type': {
						'distribution'       : df[ca.exudate_type].value_counts().to_dict(),
						'healing_correlation': df.groupby(ca.exudate_type)[hm.healing_rate].mean().to_dict()
					},
					'viscosity': {
						'distribution'       : df[ca.exudate_viscosity].value_counts().to_dict(),
						'healing_correlation': df.groupby(ca.exudate_viscosity)[hm.healing_rate].mean().to_dict()
					}
				},
				'temporal_patterns': {
					'volume_progression': df.groupby('Visit Number')[ca.exudate_volume].value_counts().to_dict(),
					'type_progression'  : df.groupby('Visit Number')[ca.exudate_type].value_counts().to_dict()
				}
			}
		}

		# Add sensor data analysis if available
		stats_data['sensor_data'] = {}

		# Temperature Analysis
		if temp.center_temp in df.columns:
			stats_data['sensor_data']['temperature'] = {
				'center_temp': {
					'overall': df[temp.center_temp].agg(['mean', 'std', 'min', 'max']).to_dict(),
					'by_healing_status': df.groupby('Healing_Status')[temp.center_temp].mean().to_dict(),
					'temporal_trend': df.groupby('Visit Number')[temp.center_temp].mean().to_dict()
				}
			}

			# Add edge and peri-wound temperatures if available
			if all(col in df.columns for col in [temp.edge_temp, temp.peri_temp]):
				stats_data['sensor_data']['temperature'].update({
					'edge_temp': {
						'overall': df[temp.edge_temp].agg(['mean', 'std', 'min', 'max']).to_dict(),
						'by_healing_status': df.groupby('Healing_Status')[temp.edge_temp].mean().to_dict()
					},
					'peri_temp': {
						'overall': df[temp.peri_temp].agg(['mean', 'std', 'min', 'max']).to_dict(),
						'by_healing_status': df.groupby('Healing_Status')[temp.peri_temp].mean().to_dict()
					},
					'gradients': {
						'center_to_edge': (df[temp.center_temp] - df[temp.edge_temp]).agg(['mean', 'std']).to_dict(),
						'center_to_peri': (df[temp.center_temp] - df[temp.peri_temp]).agg(['mean', 'std']).to_dict(),
						'by_healing_status': df.groupby('Healing_Status').apply(
							lambda x: {
								'center_to_edge': (x[temp.center_temp] - x[temp.edge_temp]).mean(),
								'center_to_peri': (x[temp.center_temp] - x[temp.peri_temp]).mean()
							}
						).to_dict()
					}
				})

		# Impedance Analysis
		if imp.highest_freq_z in df.columns:
			stats_data['sensor_data']['impedance'] = {
				'magnitude': {
					'overall': df[imp.highest_freq_z].agg(['mean', 'std', 'min', 'max']).to_dict(),
					'by_healing_status': df.groupby('Healing_Status')[imp.highest_freq_z].mean().to_dict(),
					'temporal_trend': df.groupby('Visit Number')[imp.highest_freq_z].mean().to_dict()
				}
			}

			# Add complex impedance components if available
			if all(col in df.columns for col in [imp.highest_freq_z_prime, imp.highest_freq_z_double_prime]):
				stats_data['sensor_data']['impedance'].update({
					'complex_components': {
						'real': {
							'overall': df[imp.highest_freq_z_prime].agg(['mean', 'std', 'min', 'max']).to_dict(),
							'by_healing_status': df.groupby('Healing_Status')[imp.highest_freq_z_prime].mean().to_dict()
						},
						'imaginary': {
							'overall': df[imp.highest_freq_z_double_prime].agg(['mean', 'std', 'min', 'max']).to_dict(),
							'by_healing_status': df.groupby('Healing_Status')[imp.highest_freq_z_double_prime].mean().to_dict()
						}
					}
				})

		# Oxygenation Analysis
		if oxy.oxygenation in df.columns:
			stats_data['sensor_data']['oxygenation'] = {
				'oxygenation': {
					'overall'                 : df[oxy.oxygenation].agg(['mean', 'std', 'min', 'max']).to_dict(),
					'by_healing_status'       : df.groupby('Healing_Status')[oxy.oxygenation].mean().to_dict(),
					'temporal_trend'          : df.groupby('Visit Number')[oxy.oxygenation].mean().to_dict(),
					'correlation_with_healing': df[oxy.oxygenation].corr(df[hm.healing_rate]),
					'distribution_quartiles'  : pd.qcut(df[oxy.oxygenation], q=4).value_counts().to_dict()
				}
			}

			# Add hemoglobin measurements if available
			for hb_type, col in {'hemoglobin': oxy.hemoglobin,
								'oxyhemoglobin': oxy.oxyhemoglobin,
								'deoxyhemoglobin': oxy.deoxyhemoglobin}.items():
				if col in df.columns:
					stats_data['sensor_data']['oxygenation'][hb_type] = {
						'overall'                 : df[col].agg(['mean', 'std', 'min', 'max']).to_dict(),
						'by_healing_status'       : df.groupby('Healing_Status')[col].mean().to_dict(),
						'temporal_trend'          : df.groupby('Visit Number')[col].mean().to_dict(),
						'correlation_with_healing': df[col].corr(df[hm.healing_rate])
					}

		return stats_data

	def get_processed_data(self) -> pd.DataFrame:
		"""
		Process and transform the loaded wound data to prepare it for analysis.

		This method performs several data processing steps including:
		- Cleaning column names
		- Filtering out skipped visits
		- Extracting visit number information
		- Converting and formatting dates
		- Calculating days since first visit
		- Handling wound type categorization
		- Calculating wound area when not present
		- Converting numeric columns
		- Creating derived features (temperature gradients, BMI categories)
		- Calculating healing rates and metrics for each patient
		- Estimating days to heal based on healing trajectory

		Returns:
			pd.DataFrame: Processed dataframe with additional calculated columns including
						healing rates, temperature gradients, and estimated days to heal.

		Raises:
			ValueError: If no data has been loaded prior to calling this method.

		Note:
			This method does not modify the original dataframe but returns a processed copy.
		"""

		if self.df is None:
			raise ValueError("No data available. Please load data first.")

		# Get schema columns
		pi   = self.columns.patient_identifiers
		vis  = self.columns.visit_info
		wc   = self.columns.wound_characteristics
		temp = self.columns.temperature
		oxy  = self.columns.oxygenation
		dem  = self.columns.demographics
		hm   = self.columns.healing_metrics

		# Create a copy to avoid modifying original data
		df = self.df.copy()

		# Clean column names
		df.columns = df.columns.str.strip()

		# Filter out skipped visits
		df = df[df[vis.skipped_visit] != 'Yes']

		# Extract visit number from Event Name
		df['Visit Number'] = df[pi.event_name].str.extract(r'Visit (\d+)').fillna(1).astype(int)

		# Convert and format dates
		df[vis.visit_date] = pd.to_datetime(df[vis.visit_date])

		# Calculate days since first visit for each patient
		df[vis.days_since_first_visit] = df.groupby(pi.record_id)[vis.visit_date].transform(
			lambda x: (x - x.min()).dt.days
		)

		# Handle Wound Type categorization
		if wc.wound_type in df.columns:

			# Convert to string type first to handle any existing categorical
			df[wc.wound_type] = df[wc.wound_type].astype(str)

			# Replace NaN with 'Unknown'
			df[wc.wound_type] = df[wc.wound_type].replace('nan', 'Unknown')

			# Get unique categories including 'Unknown'
			categories = sorted(df[wc.wound_type].unique())

			# Now create categorical with all possible categories
			df[wc.wound_type] = pd.Categorical(df[wc.wound_type], categories=categories)

		# Calculate wound area if not present
		if wc.wound_area not in df.columns and all(col in df.columns for col in [wc.length, wc.width]):
			df[wc.wound_area] = df[wc.length] * df[wc.width]

		# Convert numeric columns
		numeric_columns = [wc.length, wc.width, hm.healing_rate, oxy.oxygenation]
		for col in numeric_columns:
			if col in df.columns:
				df[col] = pd.to_numeric(df[col].astype(str).str.replace('%', ''), errors='coerce')

		# Create derived features
		# Temperature gradients
		if all(col in df.columns for col in [temp.center_temp, temp.edge_temp, temp.peri_temp]):
			df[temp.center_edge_gradient] = df[temp.center_temp] - df[temp.edge_temp]
			df[temp.edge_peri_gradient]   = df[temp.edge_temp] - df[temp.peri_temp]
			df[temp.total_temp_gradient]  = df[temp.center_temp] - df[temp.peri_temp]

		# BMI categories
		if dem.bmi in df.columns:
			df['BMI Category'] = pd.cut(
				df[dem.bmi],
				bins=[0, 18.5, 25, 30, 35, 100],
				labels=['Underweight', 'Normal', 'Overweight', 'Obese I', 'Obese II-III']
			)

		# Calculate healing rates
		MAX_TREATMENT_DAYS = 730  # 2 years in days
		MIN_WOUND_AREA = 0

		def calculate_patient_healing_metrics(patient_data: pd.DataFrame) -> tuple:
			"""
			Calculate healing metrics for a patient based on wound area measurements across visits.

			This function analyzes a patient's wound data over multiple visits to determine
			healing rates, improvement status, and estimated days to complete healing.

			Parameters:
			----------
			patient_data : pd.DataFrame
				DataFrame containing the patient's wound measurements across visits.
				Must contain columns for 'Visit Number' and wound area.

			Returns:
			-------
			tuple
				A tuple containing three elements:
				- healing_rates (list): List of healing rates (percentage) between consecutive visits.
				- is_improving (bool): Flag indicating if the wound is showing improvement (True) or not (False).
				- estimated_days (float or np.nan): Estimated total days until complete healing based on the current healing rate. Returns np.nan if estimation is not possible or unreliable.

			Notes:
			-----
			- Healing rate is calculated as: (previous_area - current_area) / previous_area * 100
			- If patient has fewer than 2 visits, returns ([0.0], False, np.nan)
			- The estimation uses the average healing rate of positive rates only
			- Estimation returns np.nan if:
			  * The wound is not improving
			  * Current area is below MIN_WOUND_AREA
			  * Healing rate is negative or zero
			  * Estimated days exceeds MAX_TREATMENT_DAYS
			"""

			if len(patient_data) < 2:
				return [0.0], False, np.nan

			healing_rates = []
			for i, row in patient_data.iterrows():
				if row['Visit Number'] == 1 or len(patient_data[patient_data['Visit Number'] < row['Visit Number']]) == 0:
					healing_rates.append(0)
				else:
					prev_visits = patient_data[patient_data['Visit Number'] < row['Visit Number']]
					prev_visit = prev_visits[prev_visits['Visit Number'] == prev_visits['Visit Number'].max()]

					if len(prev_visit) > 0 and wc.wound_area in patient_data.columns:
						prev_area = prev_visit[wc.wound_area].values[0]
						curr_area = row[wc.wound_area]

						if prev_area > 0 and not pd.isna(prev_area) and not pd.isna(curr_area):
							healing_rate = (prev_area - curr_area) / prev_area * 100
							healing_rates.append(healing_rate)
						else:
							healing_rates.append(0)
					else:
						healing_rates.append(0)

			valid_rates = [rate for rate in healing_rates if rate > 0]
			avg_healing_rate = np.mean(valid_rates) if valid_rates else 0
			is_improving = avg_healing_rate > 0

			estimated_days = np.nan
			if is_improving and len(patient_data) > 0:
				last_visit = patient_data.iloc[-1]
				current_area = last_visit[wc.wound_area]

				if current_area > MIN_WOUND_AREA and avg_healing_rate > 0:
					daily_healing_rate = (avg_healing_rate / 100) * current_area
					if daily_healing_rate > 0:
						days_to_heal = current_area / daily_healing_rate
						total_days = last_visit[vis.days_since_first_visit] + days_to_heal
						if 0 < total_days < MAX_TREATMENT_DAYS:
							estimated_days = float(total_days)

			return healing_rates, is_improving, estimated_days

		# Process each patient's data
		for patient_id in df[pi.record_id].unique():
			patient_data = df[df[pi.record_id] == patient_id].sort_values(vis.days_since_first_visit)
			healing_rates, is_improving, estimated_days = calculate_patient_healing_metrics(patient_data)

			# Update patient records with healing rates
			for i, (idx, row) in enumerate(patient_data.iterrows()):
				if i < len(healing_rates):
					df.loc[idx, hm.healing_rate] = healing_rates[i]

			# Update the last visit with overall improvement status
			df.loc[patient_data.iloc[-1].name, hm.overall_improvement] = 'Yes' if is_improving else 'No'

			if not np.isnan(estimated_days):
				df.loc[patient_data.index, hm.estimated_days_to_heal] = estimated_days

		# Calculate and store average healing rates
		df[hm.average_healing_rate] = df.groupby(pi.record_id)[hm.healing_rate].transform('mean')

		# Ensure estimated days column exists
		if hm.estimated_days_to_heal not in df.columns:
			df[hm.estimated_days_to_heal] = pd.Series(np.nan, index=df.index, dtype=float)

		return df

	def _extract_patient_metadata(self, patient_data) -> Dict:
		"""
		Extracts and formats patient metadata from the provided patient data.

		This method compiles demographic information, lifestyle factors, and medical history
		into a structured dictionary format. It handles missing values by setting them to None
		and processes medical history from both standard columns and free text fields.

		Parameters
		----------
		patient_data : pandas.Series or dict
			A row of patient data containing demographic information, lifestyle factors, and medical history.

		Returns
		-------
		Dict
			A dictionary containing the patient's metadata with the following keys:
			- Basic demographics (age, sex, race, ethnicity)
			- Physical measurements (weight, height, bmi)
			- Study information (study_cohort)
			- Lifestyle factors (smoking_status, packs_per_day, years_smoking, alcohol_use, alcohol_frequency)
			- Medical history (as a nested dictionary)
			- Diabetes information (status, hemoglobin_a1c, a1c_available)

		Notes
		-----
		The method uses column name mappings from the self.columns object to
		access the appropriate fields in the patient data.
		"""

		# Get column names from the schema
		dem = self.columns.demographics
		pi  = self.columns.patient_identifiers
		ls  = self.columns.lifestyle
		mh  = self.columns.medical_history

		metadata = {
			'age'              : patient_data[dem.age_at_enrollment] if not pd.isna(patient_data.get(dem.age_at_enrollment)) else None,
			'sex'              : patient_data[dem.sex] if not pd.isna(patient_data.get(dem.sex)) else None,
			'race'             : patient_data[dem.race] if not pd.isna(patient_data.get(dem.race)) else None,
			'ethnicity'        : patient_data[dem.ethnicity] if not pd.isna(patient_data.get(dem.ethnicity)) else None,
			'weight'           : patient_data[dem.weight] if not pd.isna(patient_data.get(dem.weight)) else None,
			'height'           : patient_data[dem.height] if not pd.isna(patient_data.get(dem.height)) else None,
			'bmi'              : patient_data[dem.bmi] if not pd.isna(patient_data.get(dem.bmi)) else None,
			'study_cohort'     : patient_data[dem.study_cohort] if not pd.isna(patient_data.get(dem.study_cohort)) else None,
			'smoking_status'   : patient_data[ls.smoking_status] if not pd.isna(patient_data.get(ls.smoking_status)) else None,
			'packs_per_day'    : patient_data[ls.packs_per_day] if not pd.isna(patient_data.get(ls.packs_per_day)) else None,
			'years_smoking'    : patient_data[ls.years_smoked] if not pd.isna(patient_data.get(ls.years_smoked)) else None,
			'alcohol_use'      : patient_data[ls.alcohol_status] if not pd.isna(patient_data.get(ls.alcohol_status)) else None,
			'alcohol_frequency': patient_data[ls.alcohol_drinks] if not pd.isna(patient_data.get(ls.alcohol_drinks)) else None
		}

		# Medical history from individual columns
		medical_conditions = [
			mh.respiratory, mh.cardiovascular, mh.gastrointestinal, mh.musculoskeletal,
			mh.endocrine_metabolic, mh.hematopoietic, mh.hepatic_renal, mh.neurologic, mh.immune
		]

		# Get medical history from standard columns
		metadata['medical_history'] = {
			condition: patient_data[condition]
			for condition in medical_conditions if not pd.isna(patient_data.get(condition))
		}

		# Check additional medical history from free text field
		other_history = patient_data.get(mh.medical_history)
		if not pd.isna(other_history):
			existing_conditions = set(medical_conditions)
			other_conditions = [cond.strip() for cond in str(other_history).split(',')]
			other_conditions = [cond for cond in other_conditions if cond and cond not in existing_conditions]
			if other_conditions:
				metadata['medical_history']['other'] = ', '.join(other_conditions)

		# Diabetes information
		metadata['diabetes'] = {
			'status': patient_data.get(mh.diabetes),
			'hemoglobin_a1c': patient_data.get(mh.a1c),
			'a1c_available': patient_data.get(mh.a1c_available)
		}

		return metadata

	def _get_wound_info(self, visit_data) -> Dict:
		"""
		Extract and structure detailed wound information from a single visit record.

		This method processes raw visit data to extract wound characteristics and clinical assessment
		information into a structured dictionary format. It handles missing data by converting NaN values
		to None.

		Args:
			visit_data (Dict): A dictionary containing wound data from a single visit

		Returns:
			Dict: Structured wound information with the following keys:
				- location: Anatomical location of the wound
				- type: Classification of wound type
				- current_care: Current wound care regimen
				- clinical_events: Notable clinical events
				- undermining: Dictionary containing undermining presence, location and tunneling details
				- infection: Dictionary with infection status and WiFi classification
				- granulation: Dictionary with tissue coverage and quality metrics
				- necrosis: Necrotic tissue assessment
				- exudate: Dictionary containing volume, viscosity and type of wound drainage

		Raises:
			Exception: Logs a warning and returns an empty dictionary if data processing fails
		"""
		try:
			# Get column schema for wound characteristics and clinical assessment
			wc = self.columns.wound_characteristics
			ca = self.columns.clinical_assessment

			def clean_field(data, field):
				return data.get(field) if not pd.isna(data.get(field)) else None

			present = clean_field(visit_data, wc.undermining)

			wound_info = {
				'location'       : clean_field(visit_data, wc.wound_location),
				'type'           : clean_field(visit_data, wc.wound_type),
				'current_care'   : clean_field(visit_data, wc.current_wound_care),
				'clinical_events': clean_field(visit_data, ca.clinical_events),
				'undermining': {
					'present'  : None if present is None else present == 'Yes',
					'location' : visit_data.get(wc.undermining_location),
					'tunneling': visit_data.get(wc.tunneling_location)
				},
				'infection': {
					'status'             : clean_field(visit_data, ca.infection),
					'wifi_classification': visit_data.get(ca.wifi_classification)
				},
				'granulation': {
					'coverage': clean_field(visit_data, ca.granulation),
					'quality' : clean_field(visit_data, ca.granulation_quality)
				},
				'necrosis': visit_data.get(ca.necrosis),
				'exudate': {
					'volume'   : visit_data.get(ca.exudate_volume),
					'viscosity': visit_data.get(ca.exudate_viscosity),
					'type'     : visit_data.get(ca.exudate_type)
				}
			}

			return wound_info

		except Exception as e:
			logger.warning(f"Error getting wound info: {str(e)}")
			return {}

	def _process_visit_data(self, visit, record_id: int) -> Optional[Dict]:
		"""
			Process the data from a single patient visit and extract relevant information.

			This method extracts and processes various measurements taken during a patient visit, including
			wound measurements, temperature readings, oxygenation data, and impedance values. It handles
			missing data gracefully by converting NaN values to None.

			Args:
				visit: A dictionary-like object containing the raw visit data
				record_id (int): The unique identifier for the patient record

			Returns:
				Optional[Dict]: A structured dictionary containing the processed visit data, or None if
				the visit date is missing. The dictionary includes:
					- visit_date: The formatted date of the visit
					- wound_measurements: Dict containing length, width, depth, and area
					- sensor_data: Dict containing:
						- oxygenation: Overall oxygenation value
						- temperature: Dict with center, edge, and peri readings
						- impedance: Dict with high, center, and low frequency measurements
						- hemoglobin: Various hemoglobin measurements

			Notes:
				The impedance data is either extracted from Excel sweep files if available,
				or from the visit parameters directly as a fallback.
		"""

		# Get column schema
		vis  = self.columns.visit_info
		wc   = self.columns.wound_characteristics
		temp = self.columns.temperature
		oxy  = self.columns.oxygenation
		imp  = self.columns.impedance

		visit_date = pd.to_datetime(visit[vis.visit_date]).strftime('%m-%d-%Y') if not pd.isna(visit.get(vis.visit_date)) else None

		if not visit_date:
			logger.warning("Missing visit date")
			return None

		def get_float(data, key):
			return float(data[key]) if not pd.isna(data.get(key)) else None

		def get_impedance_data():
			"""
			Extracts and processes impedance data for a wound assessment.

			This function retrieves electrical impedance data either from a frequency sweep file
			or from visit parameters. It processes data for three frequency points (high, center, and low)
			and calculates impedance (Z), resistance, capacitance, and frequency values for each.

			Returns:
				dict: A nested dictionary containing processed impedance data structured as:
					{
						'high_frequency': {
							'Z': float or None,             # Impedance magnitude
							'resistance': float or None,    # Real part of impedance (Z')
							'capacitance': float or None,   # Calculated from imaginary part (1/(2πf*Z''))
							'frequency': float or None      # Frequency in Hz
						},
						'center_frequency': {...},          # Same structure as high_frequency
						'low_frequency': {...}              # Same structure as high_frequency

			Notes:
				- If frequency sweep data is available, values are extracted for all three frequency points
				- If sweep data is unavailable, only high frequency data is extracted from visit parameters
				- Capacitance is calculated using the formula: C = 1/(2πf*Z'')
			"""

			impedance_data = {
				'high_frequency'  : {'Z': None, 'resistance': None, 'capacitance': None, 'frequency': None},
				'center_frequency': {'Z': None, 'resistance': None, 'capacitance': None, 'frequency': None},
				'low_frequency'   : {'Z': None, 'resistance': None, 'capacitance': None, 'frequency': None}
			}

			def transform_impedance_data(freq_data, frequency):
				if freq_data is not None:
					return {
						'Z': freq_data['Z'],
						'resistance': freq_data['Z_prime'],
						'capacitance': None if freq_data['Z_double_prime'] is None else 1 / (2 * 3.14 * frequency * freq_data['Z_double_prime']),
						'frequency': frequency
					}
				return {'Z': None, 'resistance': None, 'capacitance': None, 'frequency': None}

			df = ImpedanceAnalyzer.process_impedance_sweep_xlsx(impedance_freq_sweep_path=self.impedance_freq_sweep_path, record_id=record_id, visit_date_being_processed=visit_date)

			if df is not None:
				# Get data for this visit date
				visit_df = df[df['Visit date'] == visit_date]

				if not visit_df.empty:
					# Get data for all three frequencies using the index
					high_freq   = visit_df[visit_df.index == 'highest_freq'].iloc[0] if not visit_df[visit_df.index == 'highest_freq'].empty else None
					center_freq = visit_df[visit_df.index == 'center_freq'].iloc[0]  if not visit_df[visit_df.index == 'center_freq'].empty  else None
					low_freq    = visit_df[visit_df.index == 'lowest_freq'].iloc[0]  if not visit_df[visit_df.index == 'lowest_freq'].empty  else None

					impedance_data['high_frequency']   = transform_impedance_data(high_freq,   float(high_freq['frequency'])   if high_freq   is not None else None)
					impedance_data['center_frequency'] = transform_impedance_data(center_freq, float(center_freq['frequency']) if center_freq is not None else None)
					impedance_data['low_frequency']    = transform_impedance_data(low_freq,    float(low_freq['frequency'])    if low_freq    is not None else None)

			else:
				# Get impedance data from visit parameters if no sweep data
				high_freq = {
					'Z'             : get_float(visit, imp.highest_freq_z),
					'Z_prime'       : get_float(visit, imp.highest_freq_z_prime),
					'Z_double_prime': get_float(visit, imp.highest_freq_z_double_prime)
				}

				impedance_data['high_frequency'] = transform_impedance_data(high_freq, 80000)

			return impedance_data

		wound_measurements = {
			'length': get_float(visit, wc.length),
			'width' : get_float(visit, wc.width),
			'depth' : get_float(visit, wc.depth),
			'area'  : get_float(visit, wc.wound_area)
		}

		temperature_readings = {
			'center': get_float(visit, temp.center_temp),
			'edge'  : get_float(visit, temp.edge_temp),
			'peri'  : get_float(visit, temp.peri_temp)
		}

		hemoglobin_types = {
			'hemoglobin'     : oxy.hemoglobin,
			'oxyhemoglobin'  : oxy.oxyhemoglobin,
			'deoxyhemoglobin': oxy.deoxyhemoglobin
		}

		return {
			'visit_date': visit_date,
			'wound_measurements': wound_measurements,
			'sensor_data': {
				'oxygenation': get_float(visit, oxy.oxygenation),
				'temperature': temperature_readings,
				'impedance'  : get_impedance_data(),
				**{key: get_float(visit, value) for key, value in hemoglobin_types.items()}
			}
		}


@dataclass
class DataManager:
	"""Handles data loading, processing and manipulation."""
	# data_processor: WoundDataProcessor = field(default_factory=lambda: None)
	df            : pd.DataFrame       = field(default_factory=lambda: None)
	schema        : DataColumns        = field(default_factory=DataColumns)

	@staticmethod
	def load_data(csv_dataset_path):
		"""Load and preprocess the wound healing data from an uploaded CSV file. Returns None if no file is provided."""

		if csv_dataset_path is None:
			return None

		df = pd.read_csv(csv_dataset_path)
		df = DataManager._preprocess_data(df)
		return df

	@staticmethod
	def _preprocess_data(df: pd.DataFrame) -> pd.DataFrame:
		"""
		Preprocesses the wound data DataFrame.

		This function performs several preprocessing steps on the input DataFrame, including:
		1. Normalizing column names
		2. Filtering out skipped visits
		3. Extracting and converting visit numbers
		4. Formatting visit dates
		5. Converting wound type to categorical data
		6. Calculating wound area from dimensions if not present
		7. Creating derived features and healing rates

		Args:
			df (pd.DataFrame): The raw wound data DataFrame to preprocess.

		Returns:
			pd.DataFrame: The preprocessed wound data.
		"""

		# Get column names from schema
		schema = DataColumns()
		pi = schema.patient_identifiers
		vis = schema.visit_info
		wc = schema.wound_characteristics

		df.columns = df.columns.str.strip()

		# Filter out skipped visits
		df = df[df[vis.skipped_visit] != 'Yes']

		# Fill missing Visit Number with 1 before converting to int
		df['Visit Number'] = df[pi.event_name].str.extract(r'Visit (\d+)').fillna(1).astype(int)

		df[vis.visit_date] = pd.to_datetime(df[vis.visit_date]).dt.strftime('%m-%d-%Y')

		# Convert Wound Type to categorical with specified categories
		df[wc.wound_type] = pd.Categorical(df[wc.wound_type].fillna('Unknown'), categories=df[wc.wound_type].dropna().unique())

		# 2. Calculate wound area if not present but dimensions are available
		if wc.wound_area not in df.columns and all(col in df.columns for col in [wc.length, wc.width]):
			df[wc.wound_area] = df[wc.length] * df[wc.width]

		df = DataManager._create_derived_features(df)
		df = DataManager._calculate_healing_rates(df)
		return df

	@staticmethod
	def _calculate_healing_rates(df: pd.DataFrame) -> pd.DataFrame:
		"""Calculate healing rates and related metrics for each patient's wound data.

		This function processes wound measurement data to calculate:
		- Per-visit healing rates (percentage change in wound area)
		- Overall improvement status (Yes/No)
		- Average healing rate across all visits
		- Estimated days to complete healing

		The calculations are performed on a per-patient basis and consider the
		sequential nature of wound healing across multiple visits.

		Args:
			df (pd.DataFrame): DataFrame containing wound measurement data with columns defined in the DataColumns schema.

			pd.DataFrame: The input DataFrame with additional columns for healing metrics:
							- healing_rate: Per-visit healing rate
							- overall_improvement: Yes/No indicator of healing progress
							- average_healing_rate: Mean healing rate across visits
							- estimated_days_to_heal: Projected time to complete healing

		Notes:
			- Healing rate is calculated as percentage decrease in wound area between visits
			- A positive healing rate indicates improvement (wound area reduction)
			- Estimated healing time is based on the latest wound area and average healing rate
			- Calculations handle edge cases like single visits and invalid measurements
		"""

		# Get column names from schema
		schema = DataColumns()
		pi = schema.patient_identifiers
		vis = schema.visit_info
		wc = schema.wound_characteristics
		hm = schema.healing_metrics

		# Constants
		MAX_TREATMENT_DAYS = 730  # 2 years in days
		MIN_WOUND_AREA = 0

		def calculate_patient_healing_metrics(patient_data: pd.DataFrame) -> tuple[list, bool, float]:
			"""Calculate healing rate and estimated days for a patient.

			Returns:
				tuple: (healing_rates, is_improving, estimated_days_to_heal)
			"""
			if len(patient_data) < 2:
				return [0.0], False, np.nan

			# 3. Calculate healing rate (% change in wound area per visit)
			healing_rates = []
			for i, row in patient_data.iterrows():
				if row['Visit Number'] == 1 or len(patient_data[patient_data['Visit Number'] < row['Visit Number']]) == 0:
					healing_rates.append(0)  # No healing rate for first visit
				else:
					# Find the most recent previous visit
					prev_visits = patient_data[patient_data['Visit Number'] < row['Visit Number']]
					prev_visit  = prev_visits[prev_visits['Visit Number'] == prev_visits['Visit Number'].max()]

					if len(prev_visit) > 0 and wc.wound_area in patient_data.columns:
						prev_area = prev_visit[wc.wound_area].values[0]
						curr_area = row[wc.wound_area]

						# Check for valid values to avoid division by zero
						if prev_area > 0 and not pd.isna(prev_area) and not pd.isna(curr_area):
							healing_rate = (prev_area - curr_area) / prev_area * 100  # Percentage decrease
							healing_rates.append(healing_rate)
						else:
							healing_rates.append(0)
					else:
						healing_rates.append(0)

			# Calculate average healing rate and determine if improving
			valid_rates = [rate for rate in healing_rates if rate > 0]
			avg_healing_rate = np.mean(valid_rates) if valid_rates else 0
			is_improving = avg_healing_rate > 0

			# Calculate estimated days to heal based on the latest wound area and average healing rate
			estimated_days = np.nan
			if is_improving and len(patient_data) > 0:
				last_visit = patient_data.iloc[-1]
				current_area = last_visit[wc.wound_area]

				if current_area > MIN_WOUND_AREA and avg_healing_rate > 0:
					# Convert percentage rate to area change per day
					daily_healing_rate = (avg_healing_rate / 100) * current_area
					if daily_healing_rate > 0:
						days_to_heal = current_area / daily_healing_rate
						total_days = last_visit[vis.days_since_first_visit] + days_to_heal
						if 0 < total_days < MAX_TREATMENT_DAYS:
							estimated_days = float(total_days)

			return healing_rates, is_improving, estimated_days

		# Process each patient's data
		for patient_id in df[pi.record_id].unique():
			patient_data = df[df[pi.record_id] == patient_id].sort_values(vis.days_since_first_visit)

			healing_rates, is_improving, estimated_days = calculate_patient_healing_metrics(patient_data)

			# Update patient records with healing rates
			for i, (idx, row) in enumerate(patient_data.iterrows()):
				if i < len(healing_rates):
					df.loc[idx, hm.healing_rate] = healing_rates[i]

			# Update the last visit with overall improvement status
			df.loc[patient_data.iloc[-1].name, hm.overall_improvement] = 'Yes' if is_improving else 'No'

			if not np.isnan(estimated_days):
				df.loc[patient_data.index, hm.estimated_days_to_heal] = estimated_days

		# Calculate and store average healing rates
		df[hm.average_healing_rate] = df.groupby(pi.record_id)[hm.healing_rate].transform('mean')

		# Ensure estimated days column exists
		if hm.estimated_days_to_heal not in df.columns:
			df[hm.estimated_days_to_heal] = pd.Series(np.nan, index=df.index, dtype=float)

		return df

	@staticmethod
	def _create_derived_features(df: pd.DataFrame) -> pd.DataFrame:

		# Get column names from schema
		schema = DataColumns()
		pi = schema.patient_identifiers
		vis = schema.visit_info
		temp = schema.temperature
		dem = schema.demographics
		hm = schema.healing_metrics

		# Temperature gradients
		if all(col in df.columns for col in [temp.center_temp, temp.edge_temp, temp.peri_temp]):
			df[temp.center_edge_gradient] = df[temp.center_temp] - df[temp.edge_temp]
			df[temp.edge_peri_gradient] = df[temp.edge_temp] - df[temp.peri_temp]
			df[temp.total_temp_gradient] = df[temp.center_temp] - df[temp.peri_temp]

		# BMI categories
		if dem.bmi in df.columns:
			df['BMI Category'] = pd.cut(
				df[dem.bmi],
				bins=[0, 18.5, 25, 30, 35, 100],
				labels=['Underweight', 'Normal', 'Overweight', 'Obese I', 'Obese II-III']
			)

		if df is not None and not df.empty:
			# Convert Visit date to datetime if not already
			df[vis.visit_date] = pd.to_datetime(df[vis.visit_date])

			# Calculate days since first visit for each patient
			df[vis.days_since_first_visit] = df.groupby(pi.record_id)[vis.visit_date].transform(
				lambda x: (x - x.min()).dt.days
			)

			# Initialize columns with explicit dtypes
			df[hm.healing_rate] = pd.Series(0.0, index=df.index, dtype=float)
			df[hm.estimated_days_to_heal] = pd.Series(np.nan, index=df.index, dtype=float)
			df[hm.overall_improvement] = pd.Series(np.nan, index=df.index, dtype=str)

		return df

	@staticmethod
	def get_patient_data(df: pd.DataFrame, patient_id: int) -> pd.DataFrame:
		"""
		Retrieves data for a specific patient from the dataframe.

		This function filters a dataframe to return only records for a specific patient
		identified by their patient_id, and sorts the results by 'Visit Number'.

		Parameters:
			df (pd.DataFrame): The source dataframe containing patient records.
			patient_id (int): The unique identifier for the patient whose data is being retrieved.

		Returns:
			pd.DataFrame: A filtered and sorted dataframe containing only the specified patient's data.

		Note:
			This function relies on the DataColumns schema class to determine the column name for patient identifiers.
		"""

		# Get column names from schema
		schema = DataColumns()
		pi = schema.patient_identifiers

		return df[df[pi.record_id] == patient_id].sort_values('Visit Number')

	@staticmethod
	def _extract_patient_metadata(patient_data) -> Dict:
		"""
		Extract and structure patient metadata from the patient data record.

		This function extracts demographics, lifestyle factors, and medical history
		information from a patient record into a standardized dictionary format.
		Missing values (NaN) are converted to None in the returned dictionary.

		Parameters
		----------
		patient_data : pandas.Series or dict-like
			A patient data record containing demographic, lifestyle, and medical history information.
			Expected to have keys matching those defined in DataColumns schema.

		Returns
		-------
		Dict
			A structured dictionary containing:
			- Basic demographics (age, sex, race, ethnicity, weight, height, bmi, study_cohort)
			- Lifestyle factors (smoking_status, packs_per_day, years_smoking, alcohol_use, alcohol_frequency)
			- Medical history organized by body systems
			- Diabetes-specific information including status, hemoglobin A1C values, and availability

		Notes
		-----
		The function handles both standard medical history fields and additional conditions
		specified in free-text fields, combining them into a cohesive medical history record.
		"""

		# Get column names from schema
		schema = DataColumns()
		dem = schema.demographics
		ls = schema.lifestyle
		mh = schema.medical_history

		metadata = {
			'age': patient_data[dem.age_at_enrollment] if not pd.isna(patient_data.get(dem.age_at_enrollment)) else None,
			'sex': patient_data[dem.sex] if not pd.isna(patient_data.get(dem.sex)) else None,
			'race': patient_data[dem.race] if not pd.isna(patient_data.get(dem.race)) else None,
			'ethnicity': patient_data[dem.ethnicity] if not pd.isna(patient_data.get(dem.ethnicity)) else None,
			'weight': patient_data[dem.weight] if not pd.isna(patient_data.get(dem.weight)) else None,
			'height': patient_data[dem.height] if not pd.isna(patient_data.get(dem.height)) else None,
			'bmi': patient_data[dem.bmi] if not pd.isna(patient_data.get(dem.bmi)) else None,
			'study_cohort': patient_data[dem.study_cohort] if not pd.isna(patient_data.get(dem.study_cohort)) else None,
			'smoking_status': patient_data[ls.smoking_status] if not pd.isna(patient_data.get(ls.smoking_status)) else None,
			'packs_per_day': patient_data[ls.packs_per_day] if not pd.isna(patient_data.get(ls.packs_per_day)) else None,
			'years_smoking': patient_data[ls.years_smoked] if not pd.isna(patient_data.get(ls.years_smoked)) else None,
			'alcohol_use': patient_data[ls.alcohol_status] if not pd.isna(patient_data.get(ls.alcohol_status)) else None,
			'alcohol_frequency': patient_data[ls.alcohol_drinks] if not pd.isna(patient_data.get(ls.alcohol_drinks)) else None
		}

		# Medical history from individual columns
		medical_conditions = [
			mh.respiratory, mh.cardiovascular, mh.gastrointestinal, mh.musculoskeletal,
			mh.endocrine_metabolic, mh.hematopoietic, mh.hepatic_renal, mh.neurologic, mh.immune
		]

		# Get medical history from standard columns
		metadata['medical_history'] = {
			condition: patient_data[condition]
			for condition in medical_conditions if not pd.isna(patient_data.get(condition))
		}

		# Check additional medical history from free text field
		other_history = patient_data.get(mh.medical_history)
		if not pd.isna(other_history):
			existing_conditions = set(medical_conditions)
			other_conditions = [cond.strip() for cond in str(other_history).split(',')]
			other_conditions = [cond for cond in other_conditions if cond and cond not in existing_conditions]
			if other_conditions:
				metadata['medical_history']['other'] = ', '.join(other_conditions)

		# Diabetes information
		metadata['diabetes'] = {
			'status': patient_data.get(mh.diabetes),
			'hemoglobin_a1c': patient_data.get(mh.a1c),
			'a1c_available': patient_data.get(mh.a1c_available)
		}

		return metadata

	@staticmethod
	def _generate_mock_data() -> pd.DataFrame:
		"""
		Generates mock wound data for testing and development purposes.

		Creates a simulated dataset with 10 patients, each having 3 visits.
		The dataset includes wound measurements, temperature readings,
		patient characteristics, and derived healing metrics.

		Each patient starts with a random initial wound area between 10-20 units,
		which decreases by 10% with each subsequent visit.

		Returns:
			pd.DataFrame: A DataFrame containing mock patient wound data with the following columns:
				- Record ID: Patient identifier
				- Event Name: Visit description (e.g., "Visit 1")
				- Calculated Wound Area: Wound area in square units
				- Center/Edge/Peri-wound Temperature: Temperature readings in Fahrenheit
				- BMI: Body mass index
				- Diabetes?: Yes/No indicator
				- Smoking status: Never/Current/Former
				- Visit Number: Numeric visit identifier
				- Plus derived features and healing rates added by DataManager methods
		"""
		np.random.seed(42)
		n_patients = 10
		n_visits = 3
		rows = []
		for p in range(1, n_patients + 1):
			initial_area = np.random.uniform(10, 20)
			for v in range(1, n_visits + 1):
				area = initial_area * (0.9 ** (v - 1))
				rows.append({
					'Record ID': p,
					'Event Name': f'Visit {v}',
					'Calculated Wound Area': area,
					'Center of Wound Temperature (Fahrenheit)': np.random.uniform(97, 102),
					'Edge of Wound Temperature (Fahrenheit)': np.random.uniform(96, 101),
					'Peri-wound Temperature (Fahrenheit)': np.random.uniform(95, 100),
					'BMI': np.random.uniform(18, 35),
					'Diabetes?': np.random.choice(['Yes', 'No']),
					'Smoking status': np.random.choice(['Never', 'Current', 'Former'])
				})
		df = pd.DataFrame(rows)
		df['Visit Number'] = df['Event Name'].str.extract(r'Visit (\d+)').fillna(1).astype(int)
		df = DataManager._create_derived_features(df)
		df = DataManager._calculate_healing_rates(df)
		return df

	@staticmethod
	def create_and_save_report(patient_metadata: dict, analysis_results: str, report_path: str=None, prompt: dict=None) -> str:
		"""
			Create a report document from analysis results and save it.

			This function creates a Word document containing the analysis results for a patient,
			formats it using the WoundDataProcessor utility, and saves it to disk.

			Parameters
			----------
			patient_metadata : dict
				Dictionary containing patient information such as name, ID, and demographic data.
			analysis_results : str
				The analysis results text to be included in the report.
			prompt : dict, optional
				Optional dictionary containing prompt information that was used to generate the analysis.

			Returns
			-------
			str
				The file path of the saved report document.
		"""

		# Save the document
		if report_path is None:
			# Create logs directory if it doesn't exist
			log_dir = pathlib.Path(__file__).parent / 'logs'
			log_dir.mkdir(exist_ok=True)
			timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
			report_path = log_dir / f'wound_analysis_{timestamp}.docx'


		doc = Document()
		DataManager.format_word_document(doc=doc, analysis_results=analysis_results, patient_metadata=patient_metadata, report_path=report_path)

		return report_path

	@staticmethod
	def download_word_report(st, report_path: str):
		"""
			Creates a download button in a Streamlit app to download a Word document (.docx) report.

			Parameters:
			-----------
			st : streamlit
				The Streamlit module instance used to render UI components.
			report_path : str
				The file path to the Word document (.docx) to be downloaded.

			Returns:
			--------
			None

			Raises:
			-------
			Exception
				If there's an error reading the report file or creating the download button,
				an error message is displayed in the Streamlit app.

			Notes:
			------
			This function reads the specified Word document as binary data and creates
			a download button in the Streamlit interface.
		"""

		try:
			with open(report_path, 'rb') as f:
				bytes_data = f.read()
				st.download_button(
					label="Download Full Report (DOCX)",
					data=bytes_data,
					file_name=os.path.basename(report_path),
					mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
				)
		except Exception as e:
			st.error(f"Error preparing report download: {str(e)}")

	@staticmethod
	def format_word_document(doc: Document, analysis_results: str, patient_metadata: dict=None, report_path: str = None):
		"""
			Formats and saves a Microsoft Word document with wound care analysis results.

			This function creates a structured report document with patient information and
			analysis results, applying proper formatting to different sections of content.

			Parameters
			----------
			doc : Document
				A python-docx Document object to be formatted.
			analysis_results : str
				The wound care analysis results as a string, potentially containing markdown-style
				formatting like '**' for bold text and bullet points.
			patient_metadata : dict, optional
				Dictionary containing patient information with keys such as 'age', 'sex', 'bmi',
				and nested 'diabetes' information. If None, patient information section is omitted.
			report_path : str, optional
				Path where the document should be saved. If None, a default path is generated
				in a 'logs' directory with a timestamp.

			Returns
			-------
			str
				String representation of the path where the document was saved.

			Notes
			-----
			The function handles markdown-style formatting in the analysis_results:
			- Text surrounded by '**' is converted to bold
			- Lines starting with '- ' or '* ' are converted to bullet points
		"""

		# Add title
		title = doc.add_heading('Wound Care Analysis Report', 0)
		title.alignment = WD_ALIGN_PARAGRAPH.CENTER

		if patient_metadata is not None:
			# Add patient information section
			doc.add_heading('Patient Information', level=1)
			patient_info = doc.add_paragraph()
			patient_info.add_run('Patient Demographics:\n').bold = True
			patient_info.add_run(f"Age: {patient_metadata.get('age', 'Unknown')} years\n")
			patient_info.add_run(f"Sex: {patient_metadata.get('sex', 'Unknown')}\n")
			patient_info.add_run(f"BMI: {patient_metadata.get('bmi', 'Unknown')}\n")

			# Add diabetes information
			diabetes_info = doc.add_paragraph()
			diabetes_info.add_run('Diabetes Status:\n').bold = True
			if 'diabetes' in patient_metadata:
				diabetes_info.add_run(f"Type: {patient_metadata['diabetes'].get('status', 'Unknown')}\n")
				diabetes_info.add_run(f"HbA1c: {patient_metadata['diabetes'].get('hemoglobin_a1c', 'Unknown')}%\n")

		# Add analysis section
		doc.add_heading('Analysis Results', level=1)

		# Split analysis into sections and format them
		sections = analysis_results.split('\n\n')
		for section in sections:
			if section.strip():
				if '**' in section:  # Handle markdown-style headers
					# Convert markdown headers to proper formatting
					section = section.replace('**', '')
					p = doc.add_paragraph()
					p.add_run(section.strip()).bold = True
				else:
					# Handle bullet points
					if section.strip().startswith('- ') or section.strip().startswith('* '):
						p = doc.add_paragraph(section.strip()[2:], style='List Bullet')
					else:
						p = doc.add_paragraph(section.strip())

		# Add footer with timestamp
		doc.add_paragraph(f"\nReport generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

		# Save the document
		if report_path is not None:
			doc.save(report_path)


class ImpedanceAnalyzer:
	"""Handles advanced bioimpedance analysis and clinical interpretation."""

	@staticmethod
	def calculate_visit_changes(current_visit, previous_visit):
		"""
			Calculate the percentage changes in impedance parameters between two visits and determine clinical significance.

			This function compares impedance data (Z, resistance, capacitance) at different frequencies
			between the current visit and a previous visit. It calculates percentage changes and determines
			if these changes are clinically significant based on predefined thresholds.

			Parameters
			----------
			current_visit : dict
				Dictionary containing the current visit data with sensor_data.impedance measurements
				at low, center, and high frequencies.
			previous_visit : dict
				Dictionary containing the previous visit data with sensor_data.impedance measurements
				at low, center, and high frequencies.

			Returns
			-------
			tuple
				A tuple containing two dictionaries:
				- changes: Dictionary mapping parameter names (e.g., 'resistance_low_frequency') to their percentage changes between visits.
				- clinically_significant: Dictionary mapping parameter names to boolean values indicating whether the change is clinically significant.

			Notes
			-----
			The function uses the following clinical significance thresholds:
			- Resistance: 15% change
			- Capacitance: 20% change
			- Absolute impedance (Z): 15% change

			If either value is zero or invalid, the comparison for that parameter is skipped.
		"""

		changes = {}
		clinically_significant = {}

		# Define significance thresholds based on clinical literature
		significance_thresholds = {
			'resistance': 0.15,  # 15% change is clinically significant
			'capacitance': 0.20, # 20% change is clinically significant
			'Z': 0.15  # 15% change is clinically significant for absolute impedance
		}

		freq_types = ['low_frequency', 'center_frequency', 'high_frequency']
		params = ['Z', 'resistance', 'capacitance']

		for freq_type in freq_types:
			current_freq_data = current_visit.get('sensor_data', {}).get('impedance', {}).get(freq_type, {})
			previous_freq_data = previous_visit.get('sensor_data', {}).get('impedance', {}).get(freq_type, {})

			for param in params:
				try:
					current_val = float(current_freq_data.get(param, 0))
					previous_val = float(previous_freq_data.get(param, 0))

					if previous_val != 0 and current_val != 0:
						percent_change = (current_val - previous_val) / previous_val
						key = f"{param}_{freq_type}"
						changes[key] = percent_change

						# Determine clinical significance
						is_significant = abs(percent_change) > significance_thresholds.get(param, 0.15)
						clinically_significant[key] = is_significant
				except (ValueError, TypeError, ZeroDivisionError):
					continue

		return changes, clinically_significant

	@staticmethod
	def calculate_tissue_health_index(visit):
		"""
		Calculate a tissue health index based on bioimpedance measurements from a patient visit.

		This function analyzes impedance data at different frequencies to determine tissue health.
		It calculates a health score using the ratio of low frequency to high frequency impedance,
		and optionally incorporates phase angle data when available.

		Parameters:
		-----------
		visit : dict
			A dictionary containing visit data, with the following structure:
			{
				'sensor_data': {
					'impedance': {
						'low_frequency': {
							'Z': float,  # Impedance magnitude at low frequency
							'frequency': float  # Actual frequency value
						},
						'high_frequency': {
							'Z': float,  # Impedance magnitude at high frequency
							'resistance': float,  # Optional
							'capacitance': float,  # Optional
							'frequency': float  # Optional, defaults to 80000 Hz if not provided
						}
					}
				}
			}

		Returns:
		--------
		tuple(float or None, str)
			A tuple containing:
			- health_score: A normalized score from 0-100 representing tissue health, or None if calculation fails
			- interpretation: A string describing the tissue health status or an error message

		Notes:
		------
		The health score is calculated based on:
		1. Low/high frequency impedance ratio (LF/HF ratio):
			- Optimal range is 5-12, with 8.5 being ideal
			- Scores decrease as the ratio deviates from this range

		2. Phase angle (when resistance and capacitance are available):
			- Calculated as arctan(1/(2πfRC))
			- Optimal range is 5-7 degrees
			- Higher values (up to 7 degrees) indicate better tissue health

		The final score is a weighted average: 60% from the ratio score and 40% from phase angle score.
		If phase angle cannot be calculated, only the ratio score is used.
		"""

		sensor_data    = visit.get('sensor_data', {})
		impedance_data = sensor_data.get('impedance', {})

		# Extract absolute impedance at different frequencies
		low_freq  = impedance_data.get('low_frequency', {})
		high_freq = impedance_data.get('high_frequency', {})

		low_z  = low_freq.get('Z', 0)
		high_z = high_freq.get('Z', 0)

		if low_z is None or high_z is None:
			return None, "Insufficient data for tissue health calculation"

		low_z  = float(low_z)
		high_z = float(high_z)

		if low_z > 0 and high_z > 0:
			# Calculate low/high frequency ratio
			lf_hf_ratio = low_z / high_z

			# Calculate phase angle if resistance and reactance available
			phase_angle = None
			if 'resistance' in high_freq and 'capacitance' in high_freq:
				r = float(high_freq.get('resistance', 0))
				c = float(high_freq.get('capacitance', 0))
				if r > 0 and c > 0:
					# Approximate phase angle calculation
					# Using arctan(1/(2πfRC))
					f = float(high_freq.get('frequency', 80000))
					phase_angle = math.atan(1/(2 * math.pi * f * r * c)) * (180/math.pi)

			# Normalize scores to 0-100 scale
			# Typical healthy ratio range: 5-12
			# This formula is based on bioimpedance analysis principles in wound healing:
			# - LF/HF ratio typically ranges from 5-12 in healthy tissue
			# - Optimal ratio is around 8.5 (midpoint of healthy range)
			# - Scores decrease as ratio deviates from this range
			# Within 5-12 range: Linear scaling from 100 (at 5) to 0 (at 12)
			# Outside range: Penalty increases with distance from optimal ratio
			ratio_score = max(0, min(100, (1 - (lf_hf_ratio - 5) / 7) * 100)) if 5 <= lf_hf_ratio <= 12 else max(0, 50 - abs(lf_hf_ratio - 8.5) * 5)

			if phase_angle:
				# Typical healthy phase angle range: 5-7 degrees
				# Phase angle calculation logic:
				# - Typical healthy phase angle range: 5-7 degrees
				# - Linear scaling from 0 to 100 as phase angle increases from 0 to 7 degrees
				# - Values above 7 degrees are capped at 100
				# - This approach assumes that higher phase angles (up to 7 degrees) indicate better tissue health
				phase_score = max(0, min(100, (phase_angle / 7) * 100))
				health_score = (ratio_score * 0.6) + (phase_score * 0.4)  # Weighted average
			else:
				health_score = ratio_score

			# Interpretation
			if health_score >= 80:
				interpretation = "Excellent tissue health"
			elif health_score >= 60:
				interpretation = "Good tissue health"
			elif health_score >= 40:
				interpretation = "Moderate tissue health"
			elif health_score >= 20:
				interpretation = "Poor tissue health"
			else:
				interpretation = "Very poor tissue health"

			return health_score, interpretation

		# except (ValueError, TypeError, ZeroDivisionError):
		# 	pass
		return None, "Insufficient data for tissue health calculation"

	@staticmethod
	def analyze_healing_trajectory(visits):
		"""
			Analyzes the wound healing trajectory based on impedance data from multiple visits.

			This function performs a linear regression analysis on the high-frequency impedance values
			over time to determine if there's a significant trend indicating wound healing or deterioration.

			Parameters
			-----------
			visits : list
				List of visit dictionaries, each containing visit data including sensor readings.
				Each visit dictionary should have:
				- 'visit_date': date of the visit
				- 'sensor_data': dict containing 'impedance' data with 'high_frequency' values including a 'Z' value representing impedance measurement

			Returns
			--------
			dict
				A dictionary containing:
				- 'status': 'insufficient_data' if fewer than 3 valid measurements, 'analyzed' otherwise
				- 'slope': slope of the linear regression (trend direction)
				- 'r_squared': coefficient of determination (strength of linear relationship)
				- 'p_value': statistical significance of the slope
				- 'dates': list of dates with valid measurements
				- 'values': list of impedance values used in analysis
				- 'interpretation': Clinical interpretation of results as one of:
					- "Strong evidence of healing progression"
					- "Moderate evidence of healing progression"
					- "Potential deterioration detected"
					- "No significant trend detected"

			Notes:
			------
			Negative slopes indicate healing (decreasing impedance), while positive slopes
			may indicate deterioration. The function requires at least 3 valid impedance
			readings to perform analysis.
		"""

		if len(visits) < 3:
			return {"status": "insufficient_data"}

		dates, z_values = [], []
		for visit in visits:
			try:
				high_freq = visit.get('sensor_data', {}).get('impedance', {}).get('high_frequency', {})
				z_val = float(high_freq.get('Z', 0))
				if z_val > 0:
					z_values.append(z_val)
					dates.append(visit.get('visit_date'))
			except (ValueError, TypeError):
				continue

		if len(z_values) < 3:
			return {"status": "insufficient_data"}

		# Convert dates to numerical values for regression
		x = np.arange(len(z_values))
		y = np.array(z_values)

		# Perform linear regression
		slope, intercept, r_value, p_value, std_err = stats.linregress(x, y)
		r_squared = r_value ** 2

		# Determine clinical significance of slope
		result = {
			"slope": slope,
			"r_squared": r_squared,
			"p_value": p_value,
			"dates": dates,
			"values": z_values,
			"status": "analyzed"
		}

		# Interpret the slope
		if slope < -0.5 and p_value < 0.05:
			result["interpretation"] = "Strong evidence of healing progression"
		elif slope < -0.2 and p_value < 0.10:
			result["interpretation"] = "Moderate evidence of healing progression"
		elif slope > 0.5 and p_value < 0.05:
			result["interpretation"] = "Potential deterioration detected"
		else:
			result["interpretation"] = "No significant trend detected"

		return result

	@staticmethod
	def analyze_frequency_response(visit):
		"""
		Analyze the pattern of impedance across different frequencies to assess tissue composition.

		This function utilizes bioelectrical impedance analysis (BIA) principles to evaluate
		tissue characteristics based on the frequency-dependent response to electrical current.
		It focuses on two key dispersion regions:

		1. Alpha Dispersion (low to center frequency):
			- Occurs in the kHz range
			- Reflects extracellular fluid content and cell membrane permeability
			- Large alpha dispersion may indicate edema or inflammation
			- Formula: alpha_dispersion = (Z_low - Z_center) / Z_low

		2. Beta Dispersion (center to high frequency):
			- Occurs in the MHz range
			- Reflects cellular density and intracellular properties
			- Large beta dispersion may indicate increased cellular content or structural changes
			- Formula: beta_dispersion = (Z_center - Z_high) / Z_center

		The function calculates these dispersion ratios and interprets them to assess tissue
		composition, providing insights into potential edema, cellular density, or active
		tissue remodeling.

		The interpretation logic:
		- If alpha_dispersion > 0.4 and beta_dispersion < 0.2:
			"High extracellular fluid content, possible edema"
		- If alpha_dispersion < 0.2 and beta_dispersion > 0.3:
			"High cellular density, possible granulation"
		- If alpha_dispersion > 0.3 and beta_dispersion > 0.3:
			"Mixed tissue composition, active remodeling"
		- Otherwise:
			"Normal tissue composition pattern"

		Args:
			visit (dict): A dictionary containing visit data, including sensor readings.

		Returns:
		--------
		dict
			A dictionary containing dispersion characteristics and interpretation of tissue composition.
		"""
		results = {}
		sensor_data    = visit.get('sensor_data', {})
		impedance_data = sensor_data.get('impedance', {})

		# Extract absolute impedance at different frequencies
		low_freq    = impedance_data.get('low_frequency', {})
		center_freq = impedance_data.get('center_frequency', {})
		high_freq   = impedance_data.get('high_frequency', {})

		try:
			# Get impedance values
			z_low    = low_freq.get('Z')
			z_center = center_freq.get('Z')
			z_high   = high_freq.get('Z')

			# First check if all values are not None
			if z_low is not None and z_center is not None and z_high is not None:
				# Then convert to float
				z_low    = float(z_low)
				z_center = float(z_center)
				z_high   = float(z_high)

				if z_low > 0 and z_center > 0 and z_high > 0:
					# Calculate alpha dispersion (low-to-center frequency drop)
					alpha_dispersion = (z_low - z_center) / z_low

					# Calculate beta dispersion (center-to-high frequency drop)
					beta_dispersion = (z_center - z_high) / z_center

					results['alpha_dispersion'] = alpha_dispersion
					results['beta_dispersion']  = beta_dispersion

					# Interpret dispersion patterns
					if alpha_dispersion > 0.4 and beta_dispersion < 0.2:
						results['interpretation'] = "High extracellular fluid content, possible edema"
					elif alpha_dispersion < 0.2 and beta_dispersion > 0.3:
						results['interpretation'] = "High cellular density, possible granulation"
					elif alpha_dispersion > 0.3 and beta_dispersion > 0.3:
						results['interpretation'] = "Mixed tissue composition, active remodeling"
					else:
						results['interpretation'] = "Normal tissue composition pattern"
				else:
					results['interpretation'] = "Insufficient frequency data for analysis (zero values)"
			else:
				results['interpretation'] = "Insufficient frequency data for analysis (missing values)"
		except (ValueError, TypeError, ZeroDivisionError) as e:
			error_message = f"Error processing frequency response data: {type(e).__name__}: {str(e)}"
			print(error_message)  # For console debugging
			traceback.print_exc()  # Print the full traceback
			results['interpretation'] = error_message  # Or keep the generic message if preferred

		return results

	@staticmethod
	def detect_impedance_anomalies(previous_visits, current_visit, z_score_threshold=2.5):
		"""
		Detects anomalies in impedance measurements by comparing the current visit's
		values with historical data from previous visits.

		This function analyzes three key impedance parameters:
		1. High-frequency impedance (Z)
		2. Low-frequency resistance
		3. High-frequency capacitance

		It calculates z-scores for each parameter based on historical means and standard
		deviations. Values exceeding the specified z-score threshold trigger alerts
		with clinical interpretations of the observed changes.

		Parameters:
		-----------
		previous_visits : list
			List of dictionaries containing data from previous visits, each with a
			'sensor_data' field containing impedance measurements.
		current_visit : dict
			Dictionary containing the current visit data with a 'sensor_data' field
			containing impedance measurements.
		z_score_threshold : float, optional (default=2.5)
			Threshold for flagging anomalies. Values with absolute z-scores exceeding
			this threshold will be reported.

		Returns:
		--------
		dict
			A dictionary of alerts where keys are parameter identifiers and values are
			dictionaries containing:
			- 'parameter': Display name of the parameter
			- 'z_score': Calculated z-score for the current value
			- 'direction': Whether the anomaly is an 'increase' or 'decrease'
			- 'interpretation': Clinical interpretation of the observed change

		Notes:
		------
		- Requires at least 3 previous visits with valid measurements to establish baseline
		- Ignores parameters with missing or non-positive values
		- Provides different clinical interpretations based on parameter type and direction of change
		"""

		if len(previous_visits) < 3:
			return {}

		alerts = {}

		# Parameters to monitor
		params = [
			('high_frequency', 'Z', 'High-frequency impedance'),
			('low_frequency', 'resistance', 'Low-frequency resistance'),
			('high_frequency', 'capacitance', 'High-frequency capacitance')
		]

		current_impedance = current_visit.get('sensor_data', {}).get('impedance', {})

		for freq_type, param_name, display_name in params:
			# Collect historical values
			historical_values = []

			for visit in previous_visits:
				visit_impedance = visit.get('sensor_data', {}).get('impedance', {})
				freq_data = visit_impedance.get(freq_type, {})
				try:
					value = float(freq_data.get(param_name, 0))
					if value > 0:
						historical_values.append(value)
				except (ValueError, TypeError):
					continue

			if len(historical_values) >= 3:
				# Calculate historical statistics
				mean = np.mean(historical_values)
				std = np.std(historical_values)

				# Get current value
				current_freq_data = current_impedance.get(freq_type, {})
				try:
					current_value = float(current_freq_data.get(param_name, 0))
					if current_value > 0 and std > 0:
						z_score = (current_value - mean) / std

						if abs(z_score) > z_score_threshold:
							direction = "increase" if z_score > 0 else "decrease"

							# Clinical interpretation
							if freq_type == 'high_frequency' and param_name == 'Z':
								if direction == 'increase':
									clinical_meaning = "Possible deterioration in tissue quality or increased inflammation"
								else:
									clinical_meaning = "Possible improvement in cellular integrity or reduction in edema"
							elif freq_type == 'low_frequency' and param_name == 'resistance':
								if direction == 'increase':
									clinical_meaning = "Possible decrease in extracellular fluid or improved barrier function"
								else:
									clinical_meaning = "Possible increase in extracellular fluid or breakdown of tissue barriers"
							elif freq_type == 'high_frequency' and param_name == 'capacitance':
								if direction == 'increase':
									clinical_meaning = "Possible increase in cellular density or membrane integrity"
								else:
									clinical_meaning = "Possible decrease in viable cell count or membrane dysfunction"
							else:
								clinical_meaning = "Significant change detected, clinical correlation advised"

							key = f"{freq_type}_{param_name}"
							alerts[key] = {
								"parameter": display_name,
								"z_score": z_score,
								"direction": direction,
								"interpretation": clinical_meaning
							}
				except (ValueError, TypeError):
					continue

		return alerts

	@staticmethod
	def assess_infection_risk(current_visit, previous_visit=None):
		"""
			Evaluates the infection risk for a wound based on bioimpedance measurements.

			This function analyzes bioelectrical impedance data from the current visit and optionally
			compares it with a previous visit to determine infection risk. It considers multiple factors
			including impedance ratios, changes in resistance, and phase angle calculations.

			Parameters
			----------
			current_visit : dict
				Dictionary containing current visit data with sensor_data.impedance measurements
				(including low_frequency and high_frequency values with Z, resistance, capacitance)
			previous_visit : dict, optional
				Dictionary containing previous visit data in the same format as current_visit,
				used for trend analysis

			Returns
			-------
			dict
				A dictionary containing:
				- risk_score: numeric score from 0-100 indicating infection risk
				- risk_level: string interpretation ("Low", "Moderate", or "High" infection risk)
				- contributing_factors: list of specific factors that contributed to the risk assessment

			Notes
			-----
			The function evaluates three primary factors:
			1. Low/high frequency impedance ratio (values >15 indicate increased risk)
			2. Sudden increase in low-frequency resistance
			3. Phase angle measurement

			The final risk score is capped between 0 and 100.
		"""


		risk_score = 0
		factors = []

		current_impedance = current_visit.get('sensor_data', {}).get('impedance', {})

		# Factor 1: Low/high frequency impedance ratio
		low_freq = current_impedance.get('low_frequency', {})
		high_freq = current_impedance.get('high_frequency', {})

		try:
			low_z = float(low_freq.get('Z', 0))
			high_z = float(high_freq.get('Z', 0))

			if low_z > 0 and high_z > 0:
				ratio = low_z / high_z
				# Ratios > 15 are associated with increased infection risk in literature
				if ratio > 20:
					risk_score += 40
					factors.append("Very high low/high frequency impedance ratio")
				elif ratio > 15:
					risk_score += 25
					factors.append("Elevated low/high frequency impedance ratio")
		except (ValueError, TypeError, ZeroDivisionError):
			pass

		# Factor 2: Sudden increase in low-frequency resistance (inflammatory response)
		if previous_visit:
			prev_impedance = previous_visit.get('sensor_data', {}).get('impedance', {})
			prev_low_freq = prev_impedance.get('low_frequency', {})

			try:
				current_r = float(low_freq.get('resistance', 0))
				prev_r = float(prev_low_freq.get('resistance', 0))

				if prev_r > 0 and current_r > 0:
					pct_change = (current_r - prev_r) / prev_r
					if pct_change > 0.30:  # >30% increase
						risk_score += 30
						factors.append("Significant increase in low-frequency resistance")
			except (ValueError, TypeError, ZeroDivisionError):
				pass

		# Factor 3: Phase angle calculation (if resistance and capacitance available)
		try:
			r = float(high_freq.get('resistance', 0))
			c = float(high_freq.get('capacitance', 0))
			f = float(high_freq.get('frequency', 80000))

			if r > 0 and c > 0:
				# Phase angle calculation based on the complex impedance model
				# It represents the phase difference between voltage and current in AC circuits
				# Lower phase angles indicate less healthy or more damaged tissue
				phase_angle = math.atan(1/(2 * math.pi * f * r * c)) * (180/math.pi)

				# Phase angle thresholds based on bioimpedance literature:
				# <2°: Indicates severe tissue damage or very poor health
				# 2-3°: Suggests compromised tissue health
				# >3°: Generally associated with healthier tissue
				if phase_angle < 2:
					risk_score += 30
					factors.append("Very low phase angle (<2°): Indicates severe tissue damage")
				elif phase_angle < 3:
					risk_score += 15
					factors.append("Low phase angle (2-3°): Suggests compromised tissue health")
		except (ValueError, TypeError, ZeroDivisionError):
			pass

		# Limit score to 0-100 range
		risk_score = min(100, max(0, risk_score))

		# Interpret risk level
		if risk_score >= 60:
			interpretation = "High infection risk"
		elif risk_score >= 30:
			interpretation = "Moderate infection risk"
		else:
			interpretation = "Low infection risk"

		return {
			"risk_score": risk_score,
			"risk_level": interpretation,
			"contributing_factors": factors
		}

	@staticmethod
	def calculate_cole_parameters(visit):
		"""
		Calculate Cole-Cole parameters from impedance measurement data in a visit.

		This function extracts impedance data at low, center, and high frequencies from a
		visit dictionary and calculates key Cole-Cole model parameters including:
		- R0 (low frequency resistance)
		- Rinf (high frequency resistance)
		- Cm (membrane capacitance)
		- Tau (time constant)
		- Alpha (tissue heterogeneity index)
		- Tissue homogeneity interpretation

		Parameters
		----------
		visit : dict
			A dictionary containing visit data, with a nested 'sensor_data' dictionary
			that includes impedance measurements at different frequencies

		Returns
		-------
		dict
			A dictionary containing calculated Cole-Cole parameters and tissue homogeneity assessment.
			May be empty if required data is missing or calculations fail.

		Notes
		-----
		The function handles exceptions for missing data, type errors, value errors, and
		division by zero, returning whatever parameters were successfully calculated before
		the exception occurred.
		"""

		results = {}
		impedance_data = visit.get('sensor_data', {}).get('impedance', {})

		# Extract resistance at different frequencies
		low_freq = impedance_data.get('low_frequency', {})
		center_freq = impedance_data.get('center_frequency', {})
		high_freq = impedance_data.get('high_frequency', {})

		try:
			# Get resistance values
			r_low    = float(low_freq.get('resistance', 0))
			r_center = float(center_freq.get('resistance', 0))
			r_high   = float(high_freq.get('resistance', 0))

			# Get capacitance values
			c_low    = float(low_freq.get('capacitance', 0))
			c_center = float(center_freq.get('capacitance', 0))
			c_high   = float(high_freq.get('capacitance', 0))

			# Get frequency values
			f_low    = float(low_freq.get('frequency', 100))
			f_center = float(center_freq.get('frequency', 7499))
			f_high   = float(high_freq.get('frequency', 80000))

			if r_low > 0 and r_high > 0:
				# Approximate R0 and R∞
				results['R0'] = r_low  # Low frequency resistance approximates R0
				results['Rinf'] = r_high  # High frequency resistance approximates R∞

				# Calculate membrane capacitance (Cm)
				if r_center > 0 and c_center > 0:
					results['Fc'] = f_center

					# Calculate time constant
					tau = 1 / (2 * math.pi * f_center)
					results['Tau'] = tau

					# Membrane capacitance estimation
					if (r_low - r_high) > 0 and r_high > 0:
						cm = tau / ((r_low - r_high) * r_high)
						results['Cm'] = cm

				# Calculate alpha (tissue heterogeneity)
				# Using resistance values to estimate alpha
				if r_low > 0 and r_center > 0 and r_high > 0:
					# Simplified alpha estimation
					alpha_est = 1 - (r_center / math.sqrt(r_low * r_high))
					results['Alpha'] = max(0, min(1, abs(alpha_est)))

					# Interpret alpha value
					if results['Alpha'] < 0.6:
						results['tissue_homogeneity'] = "High tissue homogeneity"
					elif results['Alpha'] < 0.8:
						results['tissue_homogeneity'] = "Moderate tissue homogeneity"
					else:
						results['tissue_homogeneity'] = "Low tissue homogeneity (heterogeneous tissue)"
		except (ValueError, TypeError, ZeroDivisionError, KeyError):
			pass

		return results

	@staticmethod
	def generate_clinical_insights(analyses):
		"""
			Generates clinical insights based on various wound analysis results.

			This function processes different aspects of wound analysis data and translates
			them into actionable clinical insights with corresponding confidence levels
			and recommendations.

			Parameters
			----------
			analyses : dict
				A dictionary containing analysis results with potential keys:
				- 'healing_trajectory': Contains slope, p_value, and status of wound healing over time
				- 'infection_risk': Contains risk_score and contributing_factors for infection
				- 'frequency_response': Contains interpretation of impedance frequency response
				- 'anomalies': Contains significant deviations in parameters with z-scores

			Returns
			-------
			list
				A list of dictionaries, each representing a clinical insight with keys:
				- 'insight': The main clinical observation
				- 'confidence': Level of certainty (High, Moderate, etc.)
				- 'recommendation': Suggested clinical action (when applicable)
				- 'supporting_factors' or 'clinical_meaning': Additional context (when available)

			Notes
			-----
			The function generates insights based on the following criteria:
			1. Healing trajectory based on impedance trends and statistical significance
			2. Infection risk assessment based on risk score
			3. Tissue composition analysis from frequency response data
			4. Anomaly detection for significant deviations in measured parameters
		"""

		insights = []

		# Healing trajectory insights
		if 'healing_trajectory' in analyses:
			trajectory = analyses['healing_trajectory']
			if trajectory.get('status') == 'analyzed':
				if trajectory.get('slope', 0) < -0.3 and trajectory.get('p_value', 1) < 0.05:
					insights.append({
						"insight": "Strong evidence of consistent wound healing progression based on impedance trends",
						"confidence": "High",
						"recommendation": "Continue current treatment protocol"
					})
				elif trajectory.get('slope', 0) > 0.3 and trajectory.get('p_value', 1) < 0.1:
					insights.append({
						"insight": "Potential stalling or deterioration in wound healing process",
						"confidence": "Moderate",
						"recommendation": "Consider reassessment of treatment approach"
					})

		# Infection risk insights
		if 'infection_risk' in analyses:
			risk = analyses['infection_risk']
			if risk.get('risk_score', 0) > 50:
				insights.append({
					"insight": f"Elevated infection risk detected ({risk.get('risk_score')}%)",
					"confidence": "Moderate to High",
					"recommendation": "Consider microbiological assessment and prophylactic measures",
					"supporting_factors": risk.get('contributing_factors', [])
				})

		# Tissue composition insights
		if 'frequency_response' in analyses:
			freq_response = analyses['frequency_response']
			if 'interpretation' in freq_response:
				insights.append({
					"insight": freq_response['interpretation'],
					"confidence": "Moderate",
					"recommendation": "Correlate with clinical assessment of wound bed"
				})

		# Anomaly detection insights
		if 'anomalies' in analyses and analyses['anomalies']:
			for param, anomaly in analyses['anomalies'].items():
				insights.append({
					"insight": f"Significant {anomaly.get('direction')} in {anomaly.get('parameter')} detected (z-score: {anomaly.get('z_score', 0):.2f})",
					"confidence": "High" if abs(anomaly.get('z_score', 0)) > 3 else "Moderate",
					"clinical_meaning": anomaly.get('interpretation', '')
				})

		return insights

	@staticmethod
	def classify_wound_healing_stage(analyses):
		"""
		Classifies the wound healing stage based on bioimpedance and tissue health analyses.

		The function determines whether the wound is in the Inflammatory, Proliferative,
		or Remodeling phase by analyzing tissue health scores, frequency response characteristics,
		and Cole parameters from bioimpedance measurements.

		Parameters:
		----------
		analyses : dict
			A dictionary containing various wound analysis results with the following keys:
			- 'tissue_health': tuple (score, description) where score is a numerical value
			- 'frequency_response': dict with keys 'alpha_dispersion' and 'beta_dispersion'
			- 'cole_parameters': dict containing Cole-Cole model parameters

		Returns:
		-------
		dict
			A dictionary with the following keys:
			- 'stage': str, the determined healing stage ('Inflammatory', 'Proliferative', or 'Remodeling')
			- 'characteristics': list, notable characteristics of the identified stage
			- 'confidence': str, confidence level of the classification ('Low', 'Moderate', or 'High')

		Notes:
		-----
		The classification uses the following general criteria:
		- Inflammatory: High alpha dispersion, low tissue health score
		- Proliferative: High beta dispersion, moderate tissue health, moderate alpha dispersion
		- Remodeling: Low alpha dispersion, high tissue health score

		If insufficient data is available, the function defaults to the Inflammatory stage with Low confidence.
		"""

		# Default to inflammatory if we don't have enough data
		stage = "Inflammatory"
		characteristics = []
		confidence = "Low"

		# Get tissue health index
		tissue_health = analyses.get('tissue_health', (None, ""))
		health_score = tissue_health[0] if tissue_health else None

		# Get frequency response
		freq_response = analyses.get('frequency_response', {})
		alpha = freq_response.get('alpha_dispersion', 0)
		beta = freq_response.get('beta_dispersion', 0)

		# Get Cole parameters
		cole_params = analyses.get('cole_parameters', {})

		# Stage classification logic
		if health_score is not None and freq_response and cole_params:
			confidence = "Moderate"

			# Inflammatory phase characteristics:
			# - High alpha dispersion (high extracellular fluid)
			# - Low tissue health score
			# - High low/high frequency ratio
			if alpha > 0.4 and health_score < 40:
				stage = "Inflammatory"
				characteristics = [
					"High extracellular fluid content",
					"Low tissue health score",
					"Elevated cellular permeability"
				]
				confidence = "High" if alpha > 0.5 and health_score < 30 else "Moderate"

			# Proliferative phase characteristics:
			# - High beta dispersion (cellular proliferation)
			# - Moderate tissue health score
			# - Moderate alpha dispersion
			elif beta > 0.3 and 40 <= health_score <= 70 and 0.2 <= alpha <= 0.4:
				stage = "Proliferative"
				characteristics = [
					"Active cellular proliferation",
					"Increasing tissue organization",
					"Moderate extracellular fluid"
				]
				confidence = "High" if beta > 0.4 and health_score > 50 else "Moderate"

			# Remodeling phase characteristics:
			# - Low alpha dispersion (reduced extracellular fluid)
			# - High tissue health score
			# - Low variability in impedance
			elif alpha < 0.2 and health_score > 70:
				stage = "Remodeling"
				characteristics = [
					"Reduced extracellular fluid",
					"Improved tissue organization",
					"Enhanced barrier function"
				]
				confidence = "High" if alpha < 0.15 and health_score > 80 else "Moderate"

		return {
			"stage": stage,
			"characteristics": characteristics,
			"confidence": confidence
		}

	def prepare_population_stats(self, df: pd.DataFrame) -> tuple:
		"""
		Calculate average impedance statistics across different groupings from the wound data.

		This function groups wound data by visit number to calculate mean impedance values,
		and also calculates the average impedance by wound type.

		Parameters
		----------
		df : pd.DataFrame
			DataFrame containing wound data with columns for 'Visit Number', 'Wound Type',
			and various impedance measurements (Z, Z', Z'').

		Returns
		-------
		tuple
			A tuple containing two DataFrames:
			- avg_impedance: DataFrame with average impedance components by visit number
			- avg_by_type: DataFrame with average impedance by wound type
		"""

		# Impedance components by visit
		avg_impedance = df.groupby('Visit Number', observed=False)[
			["Skin Impedance (kOhms) - Z",
			"Skin Impedance (kOhms) - Z'",
			"Skin Impedance (kOhms) - Z''"]
		].mean().reset_index()

		# Impedance by wound type
		avg_by_type = df.groupby('Wound Type')["Skin Impedance (kOhms) - Z"].mean().reset_index()

		return avg_impedance, avg_by_type

	def generate_clinical_analysis(self, current_visit, previous_visit=None) -> dict:
		"""
		Generate a comprehensive clinical analysis based on the current visit data and optionally compare with previous visit.

		This method analyzes wound data to produce metrics related to tissue health, infection risk,
		and frequency response. When previous visit data is provided, it also calculates changes
		between visits and identifies significant changes.

		Parameters
		----------
		current_visit : dict
			Data from the current wound assessment visit
		previous_visit : dict, optional
			Data from the previous wound assessment visit for comparison, default is None

		Returns
		-------
		dict
			A dictionary containing the following keys:
			- 'tissue_health': Index quantifying overall tissue health
			- 'infection_risk': Assessment of wound infection probability
			- 'frequency_response': Analysis of tissue composition
			- 'changes': Differences between current and previous visit (only if previous_visit provided)
			- 'significant_changes': Notable changes requiring attention (only if previous_visit provided)
		"""

		analysis = {}

		# Calculate tissue health index
		analysis['tissue_health'] = self.calculate_tissue_health_index(visit=current_visit)

		# Assess infection risk
		analysis['infection_risk'] = self.assess_infection_risk(current_visit=current_visit, previous_visit=previous_visit)

		# Analyze tissue composition (frequency response)
		analysis['frequency_response'] = self.analyze_frequency_response(visit=current_visit)

		# Calculate changes since previous visit
		if previous_visit:
			analysis['changes'], analysis['significant_changes'] = self.calculate_visit_changes( current_visit=current_visit, previous_visit=previous_visit )

		return analysis

	def generate_advanced_analysis(self, visits) -> dict:
		"""
		This method performs comprehensive analysis of wound data across multiple patient visits,
		including healing trajectory assessment, anomaly detection, Cole parameter calculations,
		tissue health evaluation, infection risk assessment, frequency response analysis,
		clinical insight generation, and wound healing stage classification.

			visits (list): List of visit dictionaries containing wound measurement data. At least 3 visits are required for full analysis.

			dict: Dictionary with advanced analysis results containing the following keys:
				- 'status': 'insufficient_data' if fewer than 3 visits are provided
				- 'healing_trajectory': Analysis of wound healing progression over time
				- 'anomalies': Detected impedance anomalies compared to previous visits
				- 'cole_parameters': Calculated Cole model parameters for the latest visit
				- 'tissue_health': Tissue health index from the most recent visit
				- 'infection_risk': Assessment of infection risk based on recent measurements
				- 'frequency_response': Analysis of bioimpedance frequency response
				- 'insights': Generated clinical insights based on all analyses
				- 'healing_stage': Classification of current wound healing stage

		Raises:
			None explicitly, but may propagate exceptions from called methods

		Note:
			This method requires at least 3 visits to generate a complete analysis.
			If fewer visits are provided, returns a dictionary with status 'insufficient_data'.
		"""
		if len(visits) < 3:
			return {'status': 'insufficient_data'}

		analysis = {}

		# Analyze healing trajectory
		analysis['healing_trajectory'] = self.analyze_healing_trajectory(visits)

		# Detect anomalies
		analysis['anomalies'] = self.detect_impedance_anomalies(
			visits[:-1], visits[-1]
		)

		# Calculate Cole parameters
		analysis['cole_parameters'] = self.calculate_cole_parameters(visits[-1])

		# Get tissue health from most recent visit
		analysis['tissue_health'] = self.calculate_tissue_health_index(visits[-1])

		# Get infection risk assessment
		analysis['infection_risk'] = self.assess_infection_risk(
			visits[-1], visits[-2] if len(visits) > 1 else None
		)

		# Get frequency response analysis
		analysis['frequency_response'] = self.analyze_frequency_response(visits[-1])

		# Generate clinical insights
		analysis['insights'] = self.generate_clinical_insights(analysis)

		# Classify wound healing stage
		analysis['healing_stage'] = self.classify_wound_healing_stage(analysis)

		return analysis

	@staticmethod
	def process_impedance_sweep_xlsx(impedance_freq_sweep_path: pathlib.Path, record_id: int, visit_date_being_processed) -> Optional[pd.DataFrame]:
		"""
			Process impedance sweep data from an Excel file for a specific record ID and visit date.

			This function reads the Excel file containing impedance sweep data for a given patient,
			extracts the relevant information for the specified visit date, and returns it as a
			processed DataFrame.

			Args:
				record_id (int): The unique identifier for the patient record.
				visit_date_being_processed (str): The specific visit date to process, in 'mm-dd-yyyy' format.

			Returns:
				Optional[pd.DataFrame]: A DataFrame containing the processed impedance sweep data
				for the specified visit date, or None if the file doesn't exist or processing fails.

			Raises:
				FileNotFoundError: If the Excel file for the given record_id is not found.
				ValueError: If the visit date cannot be extracted from the sheet name.
		"""

		# Find the Excel file for this record
		excel_file = impedance_freq_sweep_path / f'{record_id}.xlsx'
		if not excel_file.exists():
			print(f"Failed to find Excel file for record ID: {excel_file}")
			# st.error(f"Failed to find Excel file for record ID: {excel_file}")
			return None


		# Read all sheets from the Excel file
		xl = pd.ExcelFile(excel_file)

		# Process each sheet (visit date)
		dfs = []
		for sheet_name in xl.sheet_names:

			# Extract visit date from filename
			match = re.search(rf"{record_id}_visit_(\d+)_(\d+)-(\d+)-(\d+)", sheet_name)
			if not match:
				logger.warning(f"Could not extract visit date from filename: {sheet_name}")
				return None
			visit_number = match.group(1)
			visit_date = datetime.strptime(f"{match.group(2)}/{match.group(3)}/{match.group(4)}", "%m/%d/%Y").strftime('%m-%d-%Y')

			if visit_date != visit_date_being_processed:
				continue

			# Read the sheet
			df = pd.read_excel(excel_file, sheet_name=sheet_name, skiprows=1)

			# Clean column names and select relevant columns
			df.columns = df.columns.str.strip()

			# Get only the bottom half of the dataframe
			half_point = len(df) // 2
			df_bottom  = df.iloc[half_point:]

			# Find frequencies of interest in the bottom half
			lowest_freq  = df_bottom['freq / Hz'].min()
			highest_freq = df_bottom['freq / Hz'].max()

			# Calculate the difference between Z' and -Z" and find where it's minimum
			# df_bottom['z_diff'] = abs(df_bottom["Z' / Ohm"] - df_bottom["-Z'' / Ohm"])
			center_freq = df_bottom.loc[df_bottom['neg. Phase / °'].idxmax(), 'freq / Hz']

			# Get data for these three frequencies
			freq_data = []
			for freq_type, freq in [('lowest_freq', lowest_freq), ('center_freq', center_freq), ('highest_freq', highest_freq)]:
				row = df_bottom[df_bottom['freq / Hz'] == freq].iloc[0]
				freq_data.append({
					'Visit date'    : visit_date,
					'Visit number'  : visit_number,
					'frequency'     : str(freq),
					'Z'             : row["Z / Ohm"],
					'Z_prime'       : row["Z' / Ohm"],
					'Z_double_prime': row["-Z'' / Ohm"],
					'neg. Phase / °': row["neg. Phase / °"],
					'index'         : freq_type
				})

			dfs.extend(freq_data)

		if not dfs:
			return None

		# Create DataFrame from processed data
		df = pd.DataFrame(dfs).set_index('index')

		# Convert Visit date to datetime and format
		df['Visit date'] = pd.to_datetime(df['Visit date']).dt.strftime('%m-%d-%Y')

		return df


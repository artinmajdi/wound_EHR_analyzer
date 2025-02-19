import os
import argparse
from data_processor import WoundDataProcessor
from llm_interface import WoundAnalysisLLM
import logging
import pathlib
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create logs directory if it doesn't exist
log_dir = pathlib.Path(__file__).parent / 'logs'
log_dir.mkdir(exist_ok=True)

# Configure logging to both file and console
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
log_filename = log_dir / f'wound_analysis_{timestamp}.log'
word_filename = log_dir / f'wound_analysis_{timestamp}.docx'

# Create and configure logger
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_filename),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def format_word_document(doc: Document, analysis_text: str, patient_data: dict) -> None:
    """Format the analysis results in a professional Word document layout."""
    # Add title
    title = doc.add_heading('Wound Care Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add patient information section
    doc.add_heading('Patient Information', level=1)
    metadata = patient_data['patient_metadata']
    patient_info = doc.add_paragraph()
    patient_info.add_run('Patient Demographics:\n').bold = True
    patient_info.add_run(f"Age: {metadata.get('age', 'Unknown')} years\n")
    patient_info.add_run(f"Sex: {metadata.get('sex', 'Unknown')}\n")
    patient_info.add_run(f"BMI: {metadata.get('bmi', 'Unknown')}\n")

    # Add diabetes information
    diabetes_info = doc.add_paragraph()
    diabetes_info.add_run('Diabetes Status:\n').bold = True
    if 'diabetes' in metadata:
        diabetes_info.add_run(f"Type: {metadata['diabetes'].get('status', 'Unknown')}\n")
        diabetes_info.add_run(f"HbA1c: {metadata['diabetes'].get('hemoglobin_a1c', 'Unknown')}%\n")

    # Add analysis section
    doc.add_heading('Analysis Results', level=1)

    # Split analysis into sections and format them
    sections = analysis_text.split('\n\n')
    for section in sections:
        if section.strip():
            if '**' in section:  # Handle markdown-style headers
                # Convert markdown headers to proper formatting
                section = section.replace('**', '')
                p = doc.add_paragraph()
                p.add_run(section.strip()).bold = True
            else:
                # Handle bullet points
                if section.strip().startswith('- ') or section.strip().startswith('* '):
                    p = doc.add_paragraph(section.strip()[2:], style='List Bullet')
                else:
                    p = doc.add_paragraph(section.strip())

    # Add footer with timestamp
    doc.add_paragraph(f"\nReport generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

def main():
    parser = argparse.ArgumentParser(description='Analyze wound care data using LLMs')
    parser.add_argument('--record-id', type=int, default=3, help='Patient record ID to analyze')
    parser.add_argument('--model-name', type=str, default='ai-verde', choices=['falcon-7b-medical', 'biogpt', 'clinical-bert', 'ai-verde'], help='Specific model name to use')
    parser.add_argument('--dataset-path', type=pathlib.Path, help='path to three csv files', default=pathlib.Path('/Users/artinmajdi/Documents/GitHubs/postdoc/TDT_copilot_2/dataset'))
    parser.add_argument('--api-key', type=str, default="sk-h8JtQkCCJUOy-TAdDxCLGw", help='API key for AI Verde model')

    args = parser.parse_args()


    # curl -s -L "https://llm-api.cyverse.ai" -H "Authorization: Bearer sk-h8JtQkCCJUOy-TAdDxCLGw" -H 'Content-Type: application/json' | jq

    try:
        # Set API key for AI Verde if provided
        if args.api_key:
            os.environ["OPENAI_API_KEY"] = args.api_key

        # Initialize data processor
        processor = WoundDataProcessor(args.dataset_path)

        # Get patient data
        logger.info(f"Processing data for patient {args.record_id}")
        patient_data = processor.get_patient_visits(args.record_id)

        # Initialize LLM interface
        llm = WoundAnalysisLLM(model_name=args.model_name)

        # Get analysis
        logger.info("Analyzing patient data with LLM")
        analysis = llm.analyze_patient_data(patient_data)

        # Create and format Word document
        doc = Document()
        format_word_document(doc, analysis, patient_data)
        doc.save(word_filename)
        logger.info(f"Report saved to {word_filename}")

        # Print results to console
        print("\nWound Care Analysis Results:")
        print("-" * 50)
        print(analysis)

    except Exception as e:
        logger.error(f"Error during analysis: {str(e)}")
        raise

if __name__ == "__main__":
    main()
